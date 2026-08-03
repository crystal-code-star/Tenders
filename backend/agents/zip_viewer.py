"""
ZIP Viewer — browse, preview, and download files from DCE ZIP archives
====================================================================
Includes preview extraction (HTML for DOCX, full cell data for Excel).
No structured extraction – that's in separate modules.
"""

import os
import io
import zipfile
import mimetypes
import time
import logging
from pathlib import Path
from typing import Optional, Dict, List, Any

import requests
from fastapi import HTTPException

logger = logging.getLogger("zip_viewer")

SUPABASE_URL = os.getenv("SUPABASE_URL", "").rstrip("/")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY", "") or os.getenv("SUPABASE_KEY", "")
DCE_BUCKET = os.getenv("SUPABASE_DCE_BUCKET", "zip_files_tenders")
TENDERS_TABLE = os.getenv("TENDERS_TABLE", "tenders_3")


# ─── In-memory cache ──────────────────────────────────────
class _ZipCacheEntry:
    __slots__ = ("data", "fetched_at")
    def __init__(self, data: bytes):
        self.data = data
        self.fetched_at = time.time()

class ZipCache:
    def __init__(self, ttl_seconds: int = 600, max_entries: int = 30):
        self.ttl = ttl_seconds
        self.max_entries = max_entries
        self._store: Dict[str, _ZipCacheEntry] = {}

    def get(self, key: str) -> Optional[bytes]:
        entry = self._store.get(key)
        if not entry:
            return None
        if time.time() - entry.fetched_at > self.ttl:
            del self._store[key]
            return None
        return entry.data

    def set(self, key: str, data: bytes) -> None:
        if len(self._store) >= self.max_entries:
            oldest_key = min(self._store, key=lambda k: self._store[k].fetched_at)
            del self._store[oldest_key]
        self._store[key] = _ZipCacheEntry(data)

zip_cache = ZipCache()


# ─── Supabase helpers ─────────────────────────────────────
def _sb_headers() -> dict:
    return {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json",
    }

def get_tender_row(tender_id: str) -> dict:
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise HTTPException(500, "Supabase not configured")
    try:
        r = requests.get(
            f"{SUPABASE_URL}/rest/v1/{TENDERS_TABLE}",
            headers=_sb_headers(),
            params={"reference": f"eq.{tender_id}", "select": "reference,objet,dce_zip_url"},
            timeout=15,
        )
        r.raise_for_status()
        rows = r.json()
    except requests.RequestException as e:
        logger.error(f"Supabase query failed: {e}")
        raise HTTPException(502, f"Database error: {str(e)}")
    if not rows:
        raise HTTPException(404, f"Tender {tender_id} not found")
    return rows[0]

def fetch_zip_bytes(tender_id: str, dce_zip_url: Optional[str]) -> bytes:
    cached = zip_cache.get(tender_id)
    if cached is not None:
        return cached
    if not dce_zip_url:
        safe_ref = tender_id.replace("/", "_")
        dce_zip_url = f"{SUPABASE_URL}/storage/v1/object/public/{DCE_BUCKET}/{safe_ref}.zip"
    try:
        # Add retry logic for transient network errors
        for attempt in range(3):
            try:
                r = requests.get(dce_zip_url, timeout=30)
                break
            except (requests.ConnectionError, requests.Timeout) as e:
                if attempt == 2:
                    raise
                time.sleep(1)
        if r.status_code == 404:
            raise HTTPException(404, "ZIP not found")
        r.raise_for_status()
    except requests.RequestException as e:
        logger.error(f"ZIP fetch failed: {e}")
        raise HTTPException(502, f"Storage error: {str(e)}")
    data = r.content
    if not data or len(data) < 100:
        raise HTTPException(422, "Empty or corrupted ZIP")
    zip_cache.set(tender_id, data)
    return data

def open_zip(data: bytes) -> zipfile.ZipFile:
    try:
        return zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile:
        raise HTTPException(422, "Invalid ZIP")


# ─── File listing & helpers ──────────────────────────────
def guess_mime(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    overrides = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xlsm": "application/vnd.ms-excel.sheet.macroenabled.12",
        ".xls": "application/vnd.ms-excel",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".txt": "text/plain",
        ".csv": "text/csv",
    }
    if ext in overrides:
        return overrides[ext]
    guessed, _ = mimetypes.guess_type(filename)
    return guessed or "application/octet-stream"

def build_file_tree(namelist: List[str]) -> dict:
    root = {"name": "/", "type": "folder", "children": {}}
    for full_path in namelist:
        if full_path.endswith("/"):
            continue
        parts = [p for p in full_path.split("/") if p]
        node = root
        for i, part in enumerate(parts):
            is_file = i == len(parts) - 1
            if is_file:
                node["children"].setdefault(
                    part,
                    {"name": part, "type": "file", "path": full_path},
                )
            else:
                node = node["children"].setdefault(
                    part,
                    {"name": part, "type": "folder", "children": {}},
                )
    def to_list(node: dict) -> dict:
        if node["type"] == "file":
            return node
        children = sorted(
            node["children"].values(),
            key=lambda n: (n["type"] != "folder", n["name"].lower()),
        )
        return {**{k: v for k, v in node.items() if k != "children"}, "children": [to_list(c) for c in children]}
    return to_list(root)


# ─── Public API (file listing, preview, download) ────────

def list_files_in_dce(tender_id: str) -> dict:
    tender = get_tender_row(tender_id)
    zip_bytes = fetch_zip_bytes(tender_id, tender.get("dce_zip_url"))
    with open_zip(zip_bytes) as zf:
        infos = zf.infolist()
        tree = build_file_tree(zf.namelist())
        files_meta = [
            {
                "path": info.filename,
                "size_bytes": info.file_size,
                "compressed_bytes": info.compress_size,
                "mime_type": guess_mime(info.filename),
                "modified": "%04d-%02d-%02d %02d:%02d" % info.date_time[:5],
            }
            for info in infos
            if not info.is_dir()
        ]
    return {
        "tender_id": tender_id,
        "title": tender.get("objet"),
        "dce_zip_url": tender.get("dce_zip_url"),
        "file_count": len(files_meta),
        "tree": tree,
        "files": files_meta,
    }

def preview_file_from_dce(tender_id: str, file_path: str, download: bool = False):
    tender = get_tender_row(tender_id)
    zip_bytes = fetch_zip_bytes(tender_id, tender.get("dce_zip_url"))
    with open_zip(zip_bytes) as zf:
        if file_path not in zf.namelist():
            raise HTTPException(404, f"File '{file_path}' not found")
        content = zf.read(file_path)
    mime = guess_mime(file_path)
    filename = Path(file_path).name
    from fastapi.responses import StreamingResponse
    disposition = "attachment" if download else "inline"
    return StreamingResponse(
        io.BytesIO(content),
        media_type=mime,
        headers={
            "Content-Disposition": f'{disposition}; filename="{filename}"',
            "Content-Length": str(len(content)),
            "Cache-Control": "private, max-age=300",
        },
    )

def get_file_metadata(tender_id: str, file_path: str) -> dict:
    tender = get_tender_row(tender_id)
    zip_bytes = fetch_zip_bytes(tender_id, tender.get("dce_zip_url"))
    with open_zip(zip_bytes) as zf:
        if file_path not in zf.namelist():
            raise HTTPException(404, f"File '{file_path}' not found")
        info = zf.getinfo(file_path)
    return {
        "path": file_path,
        "name": Path(file_path).name,
        "size_bytes": info.file_size,
        "mime_type": guess_mime(file_path),
        "modified": "%04d-%02d-%02d %02d:%02d" % info.date_time[:5],
    }

def get_raw_file(tender_id: str, file_path: str, download: bool = False):
    return preview_file_from_dce(tender_id, file_path, download)


# ─── PREVIEW EXTRACTION (for /text/ endpoint) ────────────
# This produces HTML for DOCX and full cell data for Excel (no empty cells omitted).

def _resolve_fill_color(cell) -> Optional[str]:
    try:
        if cell.fill is None:
            return None
        fg = cell.fill.fgColor
        if fg is None:
            return None
        raw = str(fg.rgb) if fg.rgb is not None else None
        if not raw:
            return None
        raw = raw.strip("'\"")
        if raw in ("00000000", "FF000000", "FFFFFFFF", "00FFFFFF", "none", "None"):
            return None
        if len(raw) == 8:
            return f"#{raw[-6:]}"
        if len(raw) == 6:
            return f"#{raw}"
        return None
    except Exception:
        return None


def extract_text_from_dce(tender_id: str, file_path: str) -> dict:
    tender = get_tender_row(tender_id)
    zip_bytes = fetch_zip_bytes(tender_id, tender.get("dce_zip_url"))
    with open_zip(zip_bytes) as zf:
        if file_path not in zf.namelist():
            raise HTTPException(404, f"File '{file_path}' not found")
        file_bytes = zf.read(file_path)

    ext = Path(file_path).suffix.lower()

    # ── DOCX → HTML ──────────────────────────────────────────
    if ext == ".docx":
        try:
            import mammoth
            with io.BytesIO(file_bytes) as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html = result.value
                styled_html = f"""
                <html>
                <head><style>
                    body {{ font-family: Arial, sans-serif; padding: 20px; max-width: 1000px; margin: 0 auto; line-height: 1.6; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background: #f2f2f2; }}
                    ul, ol {{ margin: 0.5em 0; padding-left: 2em; }}
                    img {{ max-width: 100%; height: auto; }}
                </style></head>
                <body>{html}</body></html>
                """
                return {"type": "html", "content": styled_html}
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"Mammoth preview failed: {e}")

        # Fallback
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        text_parts.append(row_text)
            return {"type": "text", "content": "\n".join(text_parts)}
        except:
            return {"type": "text", "content": "Could not preview this DOCX file."}

    # ── Excel → JSON with full cell data ──────────────────────────────
    elif ext in [".xlsx", ".xlsm"]:
        try:
            import openpyxl
            wb_formulas = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=False)
            wb_values = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        except Exception as e:
            raise HTTPException(422, f"Invalid Excel file: {e}")

        sheets = []
        for ws_f, ws_v in zip(wb_formulas.worksheets, wb_values.worksheets):
            max_row = ws_f.max_row
            max_col = ws_f.max_column
            rows = []
            for r in range(1, max_row + 1):
                row = []
                for c in range(1, max_col + 1):
                    cell_f = ws_f.cell(row=r, column=c)
                    cell_v = ws_v.cell(row=r, column=c)

                    is_formula = isinstance(cell_f.value, str) and cell_f.value.startswith("=")
                    value = cell_v.value if cell_v.value is not None else cell_f.value

                    fill_color = _resolve_fill_color(cell_f)
                    bold = bool(cell_f.font and cell_f.font.bold)
                    italic = bool(cell_f.font and cell_f.font.italic)
                    align = cell_f.alignment.horizontal if cell_f.alignment else None
                    number_format = cell_f.number_format

                    row.append({
                        "value": value,
                        "formula": cell_f.value if is_formula else None,
                        "bold": bold,
                        "italic": italic,
                        "fill": fill_color,
                        "align": align,
                        "number_format": number_format,
                    })
                rows.append(row)

            merges = [str(m) for m in ws_f.merged_cells.ranges]
            sheets.append({
                "name": ws_f.title,
                "rows": rows,
                "row_count": max_row,
                "col_count": max_col,
                "merges": merges,
            })

        return {"sheets": sheets, "type": "xlsx"}

    # ── Old XLS ──────────────────────────────────────────────────────────
    elif ext == ".xls":
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=file_bytes)
        except Exception as e:
            raise HTTPException(422, f"Invalid XLS file: {e}")

        sheets = []
        for sheet_idx in range(wb.nsheets):
            sheet = wb.sheet_by_index(sheet_idx)
            rows = []
            for row_idx in range(sheet.nrows):
                row = []
                for col_idx in range(sheet.ncols):
                    cell = sheet.cell(row_idx, col_idx)
                    value = cell.value
                    if cell.ctype == xlrd.sheet.XL_CELL_DATE:
                        try:
                            date_tuple = xlrd.xldate_as_tuple(value, wb.datemode)
                            if date_tuple[0] > 0:
                                value = f"{date_tuple[0]:04d}-{date_tuple[1]:02d}-{date_tuple[2]:02d}"
                            else:
                                value = ""
                        except:
                            value = str(value)
                    elif cell.ctype == xlrd.sheet.XL_CELL_BOOLEAN:
                        value = "TRUE" if value else "FALSE"
                    elif cell.ctype == xlrd.sheet.XL_CELL_EMPTY:
                        value = ""
                    elif cell.ctype == xlrd.sheet.XL_CELL_NUMBER:
                        if isinstance(value, float) and value.is_integer():
                            value = int(value)
                    row.append({"value": value, "formula": None, "bold": False, "italic": False, "fill": None, "align": None, "number_format": None})
                rows.append(row)
            sheets.append({
                "name": sheet.name,
                "rows": rows,
                "row_count": sheet.nrows,
                "col_count": sheet.ncols,
                "merges": [],
            })
        return {"sheets": sheets, "type": "xls"}

    # ── Text files ──────────────────────────────────────────────────────────
    elif ext in [".txt", ".csv"]:
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            return {"content": text, "type": "text"}
        except:
            raise HTTPException(422, "Could not decode text file")

    else:
        raise HTTPException(400, "Unsupported file type for preview extraction")