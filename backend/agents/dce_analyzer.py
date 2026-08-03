"""
dce_analyzer.py — Extract DCE Information from ZIP Files (Pipeline v6 - Anti-Placeholder)
===========================================================================================
CrystalWater — Traitement d'eau & Refroidissement industriel
https://crystalwater.ma/

Pipeline:
1. ZIP → Extraction fichiers
2. Extraction texte par type
3. AAO FR uniquement (arabe ignoré) + RC
4. Recherche "visite des lieux" + "classe" + "qualification"
5. LLM → JSON (nettoyage anti-placeholder)
6. Fallback REGEX si LLM échoue
7. JSON → Supabase
"""

import os
import re
import io
import sys
import time
import json
import zipfile
import logging
from datetime import datetime
from typing import Optional, List, Dict, Any
from pathlib import Path

import requests
from dotenv import load_dotenv

load_dotenv()

# ─── EMOJIS ──────────────────────────────────────────────
ICON_PIPELINE = "🔄"
ICON_EXTRACT = "📦"
ICON_TEXT = "📄"
ICON_LLM = "🤖"
ICON_REGEX = "🔍"
ICON_JSON = "📊"
ICON_SAVE = "💾"
ICON_SUCCESS = "✅"
ICON_ERROR = "❌"
ICON_WARN = "⚠️"
ICON_AAO = "📢"
ICON_RC = "📜"
ICON_VISIT = "🏗️"
ICON_SKIP = "⏭️"
ICON_CLASSE = "🏷️"

# ─── CONFIGURATION ───────────────────────────────────────
SUPABASE_URL = os.getenv("SUPABASE_URL", "").rstrip("/")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY", "") or os.getenv("SUPABASE_KEY", "")
TENDERS_TABLE = os.getenv("TENDERS_TABLE", "tenders_3")

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")

logging.basicConfig(level=logging.INFO, format='%(message)s',
    handlers=[logging.StreamHandler(sys.stdout),
              logging.FileHandler(f'logs/dce_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log', encoding='utf-8')])
logger = logging.getLogger("dce")


# ═══════════════ SUPABASE ═══════════════

def _sb_headers(): return {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}", "Content-Type": "application/json"}
def _sb_get_tenders(p=None):
    if not SUPABASE_URL: return []
    try:
        r = requests.get(f"{SUPABASE_URL}/rest/v1/{TENDERS_TABLE}", headers=_sb_headers(), params=p or {}, timeout=15)
        return r.json() or []
    except: return []
def _sb_patch_tender(ref, data):
    if not SUPABASE_URL: return False
    try:
        h = _sb_headers(); h["Prefer"] = "return=minimal"
        r = requests.patch(f"{SUPABASE_URL}/rest/v1/{TENDERS_TABLE}?reference=eq.{ref}", headers=h, json=data, timeout=60)
        return r.status_code in (200, 204)
    except: return False


# ═══════════════ DOWNLOAD ═══════════════

def download_zip(url):
    try:
        r = requests.get(url, timeout=120)
        if r.status_code == 404: return None
        r.raise_for_status()
        return r.content if r.content and len(r.content) > 100 else None
    except: return None


# ═══════════════ EXTRACTION ═══════════════

def extract_zip(zb):
    files = {}
    try:
        zf = zipfile.ZipFile(io.BytesIO(zb))
        for fn in zf.namelist():
            if not fn.endswith('/'):
                try: files[fn] = zf.read(fn)
                except: pass
        zf.close()
    except: pass
    return files

def file_type(fn):
    return {'.pdf':'pdf','.docx':'docx','.doc':'doc','.xlsx':'xlsx','.xlsm':'xlsx','.xls':'xls','.csv':'csv','.txt':'txt','.rtf':'rtf','.xml':'xml','.html':'html'}.get(Path(fn).suffix.lower(), 'other')

def text_pdf(b):
    try:
        import fitz
        doc = fitz.open(stream=b, filetype="pdf")
        t = "\n".join([p.get_text("text").strip() for p in doc if p.get_text("text").strip()])
        doc.close()
        return t
    except: return ""

def text_docx(b):
    try:
        from docx import Document
        doc = Document(io.BytesIO(b))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells: parts.append(" | ".join(cells))
        return "\n".join(parts)
    except: return ""

def text_doc(b):
    t = text_docx(b)
    if len(t.strip()) > 100: return t
    try: return b.decode('latin-1', errors='ignore')
    except: return ""

def text_xlsx(b):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(b), data_only=True)
        parts = []
        for sn in wb.sheetnames[:10]:
            ws = wb[sn]; parts.append(f"--- {sn} ---")
            for row in ws.iter_rows(max_row=200, values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                rt = " | ".join(cells).strip()
                if rt: parts.append(rt)
        return "\n".join(parts)
    except:
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=b)
            parts = []
            for s in wb.sheets()[:10]:
                parts.append(f"--- {s.name} ---")
                for r in range(min(s.nrows, 200)):
                    cells = [str(s.cell_value(r, c)).strip() for c in range(s.ncols)]
                    rt = " | ".join([c for c in cells if c])
                    if rt: parts.append(rt)
            return "\n".join(parts)
        except: return ""

def text_plain(b):
    for enc in ['utf-8', 'latin-1', 'cp1252']:
        try:
            t = b.decode(enc, errors='ignore')
            if len(t.strip()) > 20: return t
        except: continue
    return ""

def extract_text(b, fn):
    ft = file_type(fn)
    ex = {'pdf': text_pdf, 'docx': text_docx, 'doc': text_doc, 'xlsx': text_xlsx, 'xls': text_xlsx}
    if ft in ex:
        try: return ex[ft](b)
        except: pass
    return text_plain(b)


# ═══════════════ IDENTIFICATION ═══════════════

def is_french_file(fn, txt):
    nl = fn.lower()
    if 'fr' in nl and 'ar' not in nl: return True
    if 'francais' in nl or 'français' in nl: return True
    if 'arabe' in nl or 'ar.' in nl or ' ar ' in nl: return False
    
    first = txt[:1000].lower()
    arabic_words = ["المملكة", "المغربية", "جلسة", "عمومية", "تقدير", "ضمانة", "مؤقتة", "درهم", "طلب", "عروض", "أظرفة", "متنافسين", "صفقات", "فتح", "الأظرفة", "إعلان", "يعلن", "عن", "رقم", "الثمن", "التقديري", "الضمان", "المبلغ"]
    french_words = ["estimation", "cautionnement", "provisoire", "ouvert", "appel d'offres", "maître d'ouvrage", "lot", "séance", "concurrents", "dossier", "téléchargeable", "portail", "est fixée", "est fixé", "s'élève", "dirhams"]
    
    ar_count = sum(1 for w in arabic_words if w in first)
    fr_count = sum(1 for w in french_words if w in first)
    
    if ar_count >= 3 and fr_count < 2: return False
    if fr_count >= 2: return True
    return True

def is_aao(fn, txt):
    nl = fn.lower()
    if any(p in nl for p in ["avis", "aao", "appel d'offres", "annonce"]): return True
    first = txt[:1500].lower()
    if "avis d'appel d'offres" in first: return True
    if "appel d'offres ouvert" in first: return True
    if "estimation" in first and "cautionnement" in first: return True
    if "cautionnement provisoire" in first: return True
    return False

def is_rc(fn, txt):
    nl = fn.lower()
    if any(p in nl for p in ["rc", "reglement", "règlement", "rcdp", "rcdg", "cdc", "cahier des charges", "cps", "cctp"]): return True
    first = txt[:1500].lower()
    if any(kw in first for kw in ["règlement", "reglement", "dossier technique", "cahier des clauses", "cahier des prescriptions"]): return True
    return False


# ═══════════════ RECHERCHE EXPLICITE ═══════════════

def find_visite_lieux_phrase(all_texts):
    visite_keywords = ["visite des lieux", "visite du site", "visite de chantier", "tenue de faire une visite", "tenu de faire une visite", "tenus de faire une visite", "tenue d'effectuer une visite", "visite obligatoire", "visite facultative", "obligation de visiter", "doit visiter", "doivent visiter"]
    for fn, txt in all_texts.items():
        txt_lower = txt.lower()
        for kw in visite_keywords:
            idx = txt_lower.find(kw)
            if idx >= 0:
                start = max(0, idx - 50)
                end = min(len(txt), idx + 200)
                phrase = txt[start:end].strip().replace('\n', ' ').replace('\r', ' ')
                logger.info(f"    {ICON_VISIT} Visite dans {fn}: \"{phrase[:150]}\"")
                return f"[{fn}] {phrase}"
    return None


def find_classe_phrases(all_texts):
    """
    Cherche les phrases contenant la classe/qualification dans TOUS les fichiers.
    """
    classe_keywords = [
        "classe", "qualification", "classification", "catégorie",
        "secteur", "certificat de qualification",
    ]
    phrases = []
    for fn, txt in all_texts.items():
        txt_lower = txt.lower()
        for kw in classe_keywords:
            idx = txt_lower.find(kw)
            if idx >= 0:
                start = max(0, idx - 100)
                end = min(len(txt), idx + 300)
                phrase = txt[start:end].strip().replace('\n', ' ').replace('\r', ' ')
                phrases.append(f"[{fn}] {phrase}")
                break  # Une seule phrase par fichier
    
    if phrases:
        logger.info(f"    {ICON_CLASSE} {len(phrases)} phrases de classe trouvées")
        return "\n---\n".join(phrases[:5])
    return None


# ═══════════════ PARSING MONNAIE ═══════════════

def parse_money(s):
    if not s: return None
    s = s.strip().replace(' ', '')
    if ',' in s and '.' in s:
        if s.rfind(',') > s.rfind('.'): s = s.replace('.', '').replace(',', '.')
        else: s = s.replace(',', '')
    elif ',' in s:
        parts = s.split(',')
        s = s.replace(',', '.') if len(parts[-1]) <= 2 else s.replace(',', '')
    try:
        v = float(s)
        return v if v > 100 else None
    except: return None


# ═══════════════ NETTOYAGE ANTI-PLACEHOLDER ═══════════════

PLACEHOLDER_VALUES = {
    "texte", "string", "valeur", "value", "exemple", "example",
    "description", "description_texte", "oui - x références",
    "oui - x références similaires", "oui/non", "nombre",
}

def is_placeholder(val):
    """Vérifie si une valeur est un placeholder du schéma JSON."""
    if not val or not isinstance(val, str):
        return False
    val_lower = val.strip().lower()
    if val_lower in PLACEHOLDER_VALUES:
        return True
    if val_lower in ["texte", "string", "valeur", "value"]:
        return True
    return False


def clean_llm_result(result):
    """Nettoie les valeurs placeholder retournées par le LLM."""
    cleaned = {}
    for field in ["estimation", "caution_provisoire", "caution_definitive",
                  "visite_lieux_obligatoire", "classe_demandee", "attestation_reference_demandee"]:
        val = result.get(field)
        
        if val is None or val == "" or val == 0:
            continue
        
        # Nettoyer les placeholders texte
        if field in ["classe_demandee", "attestation_reference_demandee"]:
            if isinstance(val, str) and is_placeholder(val):
                logger.info(f"    {ICON_WARN} Placeholder détecté pour {field}: '{val}' → ignoré")
                continue
            # Accepter seulement si c'est une vraie valeur
            if isinstance(val, str) and len(val.strip()) < 2:
                continue
        
        # Montants
        if field in ["estimation", "caution_provisoire", "caution_definitive"]:
            if isinstance(val, (int, float)) and val > 100:
                cleaned[field] = float(val)
            elif isinstance(val, str):
                try:
                    n = float(val.replace(' ', '').replace(',', '.'))
                    if n > 100:
                        cleaned[field] = n
                except:
                    pass
        elif field == "visite_lieux_obligatoire":
            if isinstance(val, bool):
                cleaned[field] = val
            elif isinstance(val, str):
                if val.lower() in ["true", "oui", "yes", "obligatoire"]:
                    cleaned[field] = True
                elif val.lower() in ["false", "non", "no", "facultative"]:
                    cleaned[field] = False
        else:
            if isinstance(val, str) and not is_placeholder(val):
                cleaned[field] = val
    
    return cleaned


# ═══════════════ FALLBACK REGEX ═══════════════

def regex_extract(all_text):
    combined = "\n".join(all_text)
    result = {}
    
    # Estimation
    est_patterns = [
        r"estimation\s+des\s+co[ûu]ts?\s+des\s+prestations\s+.*?(?:fix[ée]e?\s+[àa]\s+la\s+somme\s+de\s+)?.*?\(?([\d]{1,3}(?:[.\s]\d{3})*(?:,\d{2})?)\)?\s*(?:DH|DHS|dhs|MAD|dirhams)",
        r"estimation[\s\w]{0,60}?:?\s*\(?([\d]{1,3}(?:[.\s]\d{3})*(?:,\d{2})?)\)?\s*(?:DH|DHS|dhs)",
        r"l'estimation[\s\w]{0,40}?:?\s*\(?([\d]{1,3}(?:[.\s]\d{3})*(?:,\d{2})?)\)?",
    ]
    for pat in est_patterns:
        m = re.search(pat, combined, re.IGNORECASE)
        if m:
            v = parse_money(m.group(1))
            if v:
                result["estimation"] = v
                logger.info(f"    {ICON_REGEX} Estimation: {v:,.0f} DH")
                break
    
    # Caution
    cau_patterns = [
        r"cautionnement\s+provisoire\s+.*?(?:fix[ée]\s+[àa]\s+la\s+somme\s+de\s+)?.*?\(?([\d]{1,3}(?:[.\s]\d{3})*(?:,\d{2})?)\)?\s*(?:DH|DHS|dhs|dirhams)",
        r"cautionnement\s+provisoire[\s\w]{0,40}?:?\s*\(?([\d]{1,3}(?:[.\s]\d{3})*(?:,\d{2})?)\)?",
    ]
    for pat in cau_patterns:
        m = re.search(pat, combined, re.IGNORECASE)
        if m:
            v = parse_money(m.group(1))
            if v:
                result["caution_provisoire"] = v
                logger.info(f"    {ICON_REGEX} Caution: {v:,.0f} DH")
                break
    
    # Visite
    if re.search(r"tenu[e]?\s+(de|d')\s+(faire|effectuer)\s+une\s+visite", combined, re.IGNORECASE):
        result["visite_lieux_obligatoire"] = True
        logger.info(f"    {ICON_REGEX} Visite OBLIGATOIRE")
    
    # Classe - Pattern plus robuste
    classe_patterns = [
        r"classe\s*:?\s*(\d+|S|A|B|C|D|E|F)\b",
        r"cat[ée]gorie\s*:?\s*(\d+)\b",
        r"qualification\s+exig[ée]e\s+(\w+)",
        r"secteur\s+\w+\s+classe\s+(\d+)",
    ]
    for pat in classe_patterns:
        m = re.search(pat, combined, re.IGNORECASE)
        if m:
            val = m.group(1).strip().upper()
            if val.isdigit() or val in ['S', 'A', 'B', 'C', 'D', 'E', 'F']:
                result["classe_demandee"] = val
                logger.info(f"    {ICON_REGEX} Classe: {val}")
                break
    
    return result


# ═══════════════ LLM ═══════════════

SYSTEM_PROMPT = """Tu es un assistant qui extrait des données depuis des appels d'offres marocains.
Réponds UNIQUEMENT avec un objet JSON valide sur une seule ligne. Pas de markdown, pas de ```json```.
Si une information n'est PAS trouvée, mets null. Ne JAMAIS mettre "texte" ou "string" comme valeur.

RÈGLES POUR LES VALEURS:
- estimation: nombre uniquement (ex: 784080, 2666400.00). PAS "nombre", PAS "texte"
- caution_provisoire: nombre uniquement (ex: 15000, 45000.00). PAS "nombre", PAS "texte"
- caution_definitive: nombre ou null
- visite_lieux_obligatoire: true, false, ou null. PAS "true_false_null", PAS "texte"
- classe_demandee: la valeur exacte trouvée (ex: "5", "S", "4"). PAS "texte", PAS "string". Si non trouvé → null
- attestation_reference_demandee: description trouvée ou null. PAS "texte", PAS "description_texte"

JSON ATTENDU:
{"estimation":784080,"caution_provisoire":15000,"caution_definitive":null,"visite_lieux_obligatoire":null,"classe_demandee":"5","attestation_reference_demandee":null}"""


def llm_extract(aao_fr_texts, rc_text, visite_phrase, classe_phrases, tender_ref):
    if not GROQ_API_KEY: return {}
    
    combined = ""
    for fn, txt in aao_fr_texts[:3]:
        combined += f"\n=== {fn} ===\n{txt[:3000]}"
    
    if rc_text:
        rc_sections = []
        for kw in ["qualification", "classification", "classe", "référence", "categorie", "visite", "lieux", "secteur"]:
            idx = rc_text.lower().find(kw)
            if idx >= 0:
                rc_sections.append(rc_text[max(0,idx-100):min(len(rc_text),idx+800)])
        if rc_sections:
            combined += "\n\n=== RC (sections pertinentes) ===\n" + "\n---\n".join(rc_sections[:8])
        else:
            combined += "\n\n=== RC ===\n" + rc_text[:3000]
    
    if visite_phrase:
        combined += f"\n\n=== VISITE DES LIEUX ===\n{visite_phrase}"
    
    if classe_phrases:
        combined += f"\n\n=== CLASSE / QUALIFICATION ===\n{classe_phrases}"
    
    if len(combined) > 14000:
        combined = combined[:14000]
    
    logger.info(f"    {ICON_LLM} LLM: {len(aao_fr_texts)} AAO + RC + Visite + Classe = {len(combined):,} chars")
    
    try:
        r = requests.post(GROQ_API_URL,
            headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
            json={"model": GROQ_MODEL,
                  "messages": [{"role": "system", "content": SYSTEM_PROMPT},
                               {"role": "user", "content": f"Documents:\n{combined}\n\nJSON:"}],
                  "temperature": 0, "max_tokens": 400},
            timeout=30)
        
        if r.status_code == 200:
            content = r.json()["choices"][0]["message"]["content"].strip()
            content = content.replace("```json","").replace("```","").strip()
            m = re.search(r'\{[^{}]*\}', content, re.DOTALL)
            if m:
                try:
                    result = json.loads(m.group(0))
                    logger.info(f"    {ICON_JSON} LLM brut: {json.dumps(result, ensure_ascii=False)}")
                    return clean_llm_result(result)
                except:
                    pass
        return {}
    except Exception as e:
        logger.warning(f"    {ICON_ERROR} LLM: {e}")
        return {}


# ═══════════════ PIPELINE ═══════════════

def run_pipeline(zip_bytes, tender_ref):
    logger.info(f"\n  {ICON_PIPELINE} PIPELINE: {tender_ref}")
    logger.info(f"  {'─'*50}")
    
    files = extract_zip(zip_bytes)
    if not files: return {}
    
    logger.info(f"  {ICON_EXTRACT} {len(files)} fichiers")
    
    aao_fr = []
    aao_ar_skipped = 0
    rc_text = ""
    all_texts = {}
    all_text_list = []
    
    for fn, fb in files.items():
        txt = extract_text(fb, fn)
        if not txt.strip(): continue
        all_texts[fn] = txt
        all_text_list.append(txt)
        
        suffix = ""
        if is_aao(fn, txt):
            if is_french_file(fn, txt):
                aao_fr.append((fn, txt))
                suffix = f" {ICON_AAO} ✅ FR"
            else:
                aao_ar_skipped += 1
                suffix = f" {ICON_AAO} {ICON_SKIP} AR"
        elif is_rc(fn, txt) and not rc_text:
            rc_text = txt
            suffix = f" {ICON_RC}"
        
        logger.info(f"    {fn}: {len(txt):,} chars{suffix}")
    
    if aao_ar_skipped > 0:
        logger.info(f"    {ICON_SKIP} {aao_ar_skipped} AAO arabes ignorés")
    
    logger.info(f"\n  {ICON_AAO} AAO FR: {len(aao_fr)} | {ICON_RC} RC: {'oui' if rc_text else 'non'}")
    
    # Recherches explicites
    visite_phrase = find_visite_lieux_phrase(all_texts)
    classe_phrases = find_classe_phrases(all_texts)
    
    # Étape 1 : LLM
    result = llm_extract(aao_fr, rc_text, visite_phrase, classe_phrases, tender_ref)
    
    # Étape 2 : Fallback REGEX si LLM n'a pas trouvé assez
    found_count = sum(1 for v in result.values() if v is not None)
    if found_count < 2:
        logger.info(f"    {ICON_WARN} LLM: {found_count} champ(s) → Fallback REGEX")
        regex_result = regex_extract(all_text_list)
        for k, v in regex_result.items():
            if k not in result or result[k] is None:
                result[k] = v
    
    # Fallback visite
    if "visite_lieux_obligatoire" not in result and visite_phrase:
        if re.search(r"tenu[e]?\s+(de|d')\s+(faire|effectuer)\s+une\s+visite", visite_phrase, re.IGNORECASE):
            result["visite_lieux_obligatoire"] = True
    
    return result


# ═══════════════ SINGLE ═══════════════

def analyze_single(ref):
    print(f"\n  {ICON_PIPELINE} Pipeline: {ref}")
    tenders = _sb_get_tenders({"select":"reference,objet,dce_zip_url","reference":f"eq.{ref}","limit":"1"})
    if not tenders: return print(f"  {ICON_ERROR} Non trouvée")
    t = tenders[0]
    if not t.get("dce_zip_url"): return print(f"  {ICON_ERROR} Pas de DCE")
    
    print(f"  Objet: {(t.get('objet') or '')[:100]}")
    
    zb = download_zip(t["dce_zip_url"])
    if not zb: return print(f"  {ICON_ERROR} Download échoué")
    
    print(f"  {ICON_EXTRACT} {len(zb):,} bytes")
    
    result = run_pipeline(zb, ref)
    
    print(f"\n  {'─'*40}\n  RÉSULTATS:\n  {'─'*40}")
    for field, icon, label, unit in [
        ("estimation","💰","Estimation","DH"),
        ("caution_provisoire","🔒","Caution provisoire","DH"),
        ("caution_definitive","🔒","Caution définitive","DH"),
        ("visite_lieux_obligatoire","🏗️","Visite lieux",""),
        ("classe_demandee","📜","Classe",""),
        ("attestation_reference_demandee","📜","Références",""),
    ]:
        val = result.get(field)
        if val is not None and val != "":
            if isinstance(val, bool):
                d = "OBLIGATOIRE" if val else "Facultative"
            elif isinstance(val, (int, float)):
                d = f"{val:,.0f} {unit}"
            else:
                d = str(val)
            print(f"  {icon} {label}: {d}")
        else:
            print(f"  {icon} {label}: —")
    
    if result:
        if _sb_patch_tender(ref, result):
            print(f"\n  {ICON_SAVE} {ICON_SUCCESS} Sauvegardé ({len(result)} champs)")
        else:
            print(f"\n  {ICON_SAVE} {ICON_ERROR} Échec")
    else:
        print(f"\n  {ICON_WARN} Rien extrait")


# ═══════════════ MAIN ═══════════════

if __name__ == "__main__":
    print(f"\n{'='*60}\n  {ICON_PIPELINE} CrystalWater - DCE Pipeline v6\n  Anti-Placeholder + Classe explicite\n{'='*60}")
    ref = input("\n  Référence: ").strip()
    if ref: analyze_single(ref)