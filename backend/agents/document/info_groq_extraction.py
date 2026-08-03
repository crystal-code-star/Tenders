"""
info_groq_extraction.py — DCE Extraction — v13 FIXED
======================================================
CrystalWater — Traitement d'eau & Refroidissement industriel
https://crystalwater.ma/

v13:
- extract_dossier_technique() élargi : capture "le dossier technique comprend"
- Sauvegarde corrigée : mapping vers les colonnes de tenders_3
- Ajout chiffre_affaires dans le prompt
"""

import os
import re
import io
import sys
import json
import logging
import argparse
import tempfile
import zipfile
import subprocess
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Optional
from dotenv import load_dotenv
from colorama import init, Fore, Style
import supabase
import requests as http_requests

init(autoreset=True)

current_dir = Path(__file__).resolve().parent
sys.path.insert(0, str(current_dir))
env_path = current_dir.parent.parent / ".env"
load_dotenv(env_path)

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")
supabase_client = supabase.create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY) if SUPABASE_URL and SUPABASE_SERVICE_KEY else None

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")

logging.basicConfig(level=logging.WARNING, format="%(message)s")
logger = logging.getLogger(__name__)

ICON_AI, ICON_OK, ICON_KO, ICON_WARN = "🤖", "✅", "❌", "⚠️"
ICON_ZIP, ICON_DL, ICON_SAVE = "📦", "📥", "💾"
ICON_AVIS, ICON_RC = "📢", "📜"


# ═══════════════ SUPABASE ═══════════════

def _sb_headers(): return {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}", "Content-Type": "application/json"}
def _sb_get_tenders(p=None):
    if not SUPABASE_URL: return []
    try:
        r = http_requests.get(f"{SUPABASE_URL}/rest/v1/{os.getenv('TENDERS_TABLE', 'tenders_3')}", headers=_sb_headers(), params=p or {}, timeout=15)
        return r.json() or []
    except: return []
def _sb_patch_tender(ref, data):
    if not SUPABASE_URL: return False
    try:
        h = _sb_headers(); h["Prefer"] = "return=minimal"
        r = http_requests.patch(f"{SUPABASE_URL}/rest/v1/{os.getenv('TENDERS_TABLE', 'tenders_3')}?reference=eq.{ref}", headers=h, json=data, timeout=60)
        return r.status_code in (200, 204)
    except: return False


# ═══════════════ DOWNLOAD ═══════════════

def download_zip(url):
    try:
        r = http_requests.get(url, timeout=120)
        if r.status_code == 404: return None
        r.raise_for_status()
        return r.content if r.content and len(r.content) > 100 else None
    except: return None


# ═══════════════ TEXT EXTRACTION ═══════════════

def extract_text_from_docx(b: bytes) -> str:
    try:
        from docx import Document
        doc = Document(BytesIO(b))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells: parts.append(" | ".join(cells))
        return "\n".join(parts)
    except: return ""

def extract_text_from_doc(b: bytes) -> str:
    t = extract_text_from_docx(b)
    if len(t.strip()) > 200: return t
    try:
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(b); tmp_path = tmp.name
        out_dir = tempfile.mkdtemp()
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir', out_dir, tmp_path],
                       capture_output=True, text=True, timeout=60)
        for f in os.listdir(out_dir):
            if f.endswith('.txt'):
                with open(os.path.join(out_dir, f), 'r', encoding='utf-8', errors='ignore') as ff:
                    t = ff.read().strip()
                break
        os.unlink(tmp_path)
        for f in os.listdir(out_dir):
            try: os.unlink(os.path.join(out_dir, f))
            except: pass
        os.rmdir(out_dir)
        if len(t.strip()) > 200: return t
    except: pass
    try:
        t = b.decode('latin-1', errors='ignore')
        t = re.sub(r'[^\x20-\x7E\xA0-\xFF\n\r\t\.\,\;\:\!\?\(\)\[\]\{\}\-\+\/\@\#\$\%\^\&\*\=]', ' ', t)
        t = re.sub(r'\s+', ' ', t)
    except: pass
    return t

def extract_text_from_pdf(b: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=b, filetype="pdf")
        t = "\n".join(page.get_text("text").strip() for page in doc if page.get_text("text").strip())
        doc.close()
        if len(t.strip()) < 100:
            try:
                import pytesseract
                from PIL import Image
                doc = fitz.open(stream=b, filetype="pdf")
                ocr_parts = []
                for i, page in enumerate(doc):
                    if i >= 5: break
                    pix = page.get_pixmap(dpi=200)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ot = pytesseract.image_to_string(img, lang='fra+ara')
                    if ot.strip(): ocr_parts.append(ot.strip())
                doc.close()
                if ocr_parts: t = "\n".join(ocr_parts)
            except: pass
        return t
    except: return ""

def extract_text_from_xlsx(b: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(BytesIO(b), data_only=True)
        parts = []
        for sn in wb.sheetnames[:5]:
            ws = wb[sn]; parts.append(f"--- {sn} ---")
            for row in ws.iter_rows(max_row=100, values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                rt = " | ".join(cells).strip()
                if rt: parts.append(rt)
        return "\n".join(parts)
    except: return ""

def extract_text_from_any(filename: str, file_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == '.docx': return extract_text_from_docx(file_bytes)
    elif ext == '.doc': return extract_text_from_doc(file_bytes)
    elif ext == '.pdf': return extract_text_from_pdf(file_bytes)
    elif ext in ['.xlsx', '.xlsm', '.xls']: return extract_text_from_xlsx(file_bytes)
    else:
        try: return file_bytes.decode('utf-8', errors='ignore')
        except:
            try: return file_bytes.decode('latin-1', errors='ignore')
            except: return ""


# ═══════════════ IDENTIFICATION ═══════════════

def is_avis_fr(filename: str, text: str) -> bool:
    nl = filename.lower()
    if ('avis' in nl or 'aao' in nl) and 'fr' in nl and 'ar' not in nl: return True
    if 'arabe' in nl or ' ar ' in nl: return False
    first = text[:1500].lower()
    if "avis d'appel d'offres" in first: return True
    if "appel d'offres ouvert" in first: return True
    if "estimation" in first and "cautionnement" in first: return True
    return False

def is_rc_file(filename: str) -> bool:
    nl = filename.lower()
    return any(p in nl for p in ["rc", "reglement", "règlement", "rcdp", "rcdg", "consultation", "cps"])

def extract_dossier_technique(text: str) -> str:
    """
    🔥 CORRIGÉ : Cherche TOUTES les variantes de 'dossier technique'
    et prend PLUS de contenu (10000 caractères).
    """
    start_patterns = [
        r"le\s+dossier\s+technique\s+comprend",
        r"(?:B\s*[-–]\s*)?(?:Un\s+)?dossier\s+technique\s+(?:comprenant|comprenant\s*:)?",
        r"dossier\s+technique\s*:",
        r"certificat\s+de\s+qualification\s+et\s+de\s+classification",
        r"qualification\s+et\s+classification",
        r"pi[èe]ces?\s+justifiant\s+les\s+capacit[ée]s\s+techniques",
    ]
    
    start_idx = -1
    for pat in start_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            start_idx = m.start()
            break
    
    if start_idx == -1:
        # Fallback : chercher "classe" ou "qualification" ou "secteur"
        for kw in ["classe", "qualification", "secteur", "références similaires"]:
            idx = text.lower().find(kw)
            if idx >= 0:
                start_idx = max(0, idx - 500)
                break
    
    if start_idx == -1:
        return ""
    
    # Prendre 10000 caractères (au lieu de 4000/8000)
    section = text[start_idx:start_idx + 10000]
    
    # Couper aux sections suivantes
    end_markers = [
        r"\n\s*ARTICLE\s+\d+\s*[:\-]",
        r"\n\s*CHAPITRE\s+",
        r"\n\s*PRIX\s+(?:N°|N)\s*[\.\d]",
        r"\n\s*BORDEREAU",
        r"\n\s*OFFRE\s+FINANCI[ÈE]RE",
    ]
    
    for marker in end_markers:
        end_match = re.search(marker, section, re.IGNORECASE)
        if end_match and end_match.start() > 200:
            section = section[:end_match.start()]
            break
    
    return section.strip()


# ═══════════════ GROQ LLM ═══════════════

SYSTEM_PROMPT = """Tu es un assistant qui extrait des données depuis des MARCHÉS PUBLICS MAROCAINS.

Extrais UNIQUEMENT ces 8 informations. Réponds en JSON valide sur une seule ligne.
Si pas trouvé → null. N'invente JAMAIS.

CHAMPS:
- estimation: montant estimé des travaux en DH (nombre). Cherche dans l'AVIS.
- caution_provisoire: cautionnement provisoire en DH (nombre). Cherche dans l'AVIS.
- objet: titre du marché (texte).
- date_soumission: date limite JJ/MM/AAAA.
- classe_demandee: classe de qualification. Cherche "Classe X", "Secteur I | Classe 5 | Qualification I6". Format: "5" ou "Secteur I | Classe 5 | Qualification I6".
- attestation_reference_demandee: nombre de références similaires. Cherche "au moins X projets", "références similaires". Format: "2" ou "OUI" ou null.
- visite_lieux_obligatoire: true/false/null. Cherche "visite des lieux obligatoire". PAS "regard de visite".
- chiffre_affaires: montant minimum de CA en DH (nombre). Cherche "chiffre d'affaires", "CA", "capacité financière". null si pas trouvé.

JSON: {"estimation":null,"caution_provisoire":null,"objet":"...","date_soumission":"...","classe_demandee":"...","attestation_reference_demandee":"...","visite_lieux_obligatoire":null,"chiffre_affaires":null}"""

PLACEHOLDERS = {"texte","string","valeur","value","exemple","example","description","nombre","oui/non"}

def clean_result(result: dict) -> dict:
    cleaned = {}
    for field, val in result.items():
        if val is None or val == "" or val == 0: continue
        if field in ["classe_demandee", "attestation_reference_demandee"]:
            if isinstance(val, str) and (val.strip().lower() in PLACEHOLDERS or len(val.strip()) < 1): continue
        if field in ["estimation", "caution_provisoire", "chiffre_affaires"]:
            if isinstance(val, (int, float)) and val > 100: cleaned[field] = float(val)
            elif isinstance(val, str):
                try:
                    n = float(val.replace(' ','').replace(',','.'))
                    if n > 100: cleaned[field] = n
                except: pass
        elif field == "visite_lieux_obligatoire":
            if isinstance(val, bool): cleaned[field] = val
            elif isinstance(val, str):
                if val.lower() in ["true","oui","yes","obligatoire"]: cleaned[field] = True
                elif val.lower() in ["false","non","no","facultative"]: cleaned[field] = False
        else:
            if isinstance(val, str) and val.strip().lower() not in PLACEHOLDERS: cleaned[field] = val
    return cleaned

def extract_with_groq(combined_text: str, tender_ref: str) -> dict:
    if not GROQ_API_KEY or not combined_text.strip(): return {}
    if len(combined_text) > 15000:
        combined_text = combined_text[:10000] + "\n...[tronqué]...\n" + combined_text[-5000:]
    
    try:
        r = http_requests.post(GROQ_API_URL,
            headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
            json={"model": GROQ_MODEL,
                  "messages": [{"role": "system", "content": SYSTEM_PROMPT},
                               {"role": "user", "content": f"Réf: {tender_ref}\n\n{combined_text}"}],
                  "temperature": 0, "max_tokens": 500},
            timeout=30)
        if r.status_code == 200:
            content = r.json()["choices"][0]["message"]["content"].strip()
            content = content.replace("```json","").replace("```","").strip()
            m = re.search(r'\{[^{}]*\}', content, re.DOTALL)
            if m:
                try: return clean_result(json.loads(m.group(0)))
                except: pass
        return {}
    except Exception as e:
        logger.warning(f"Groq error: {e}")
        return {}


# ═══════════════ PIPELINE ═══════════════

def process_dce(zip_bytes: bytes, tender_ref: str) -> dict:
    all_files = {}
    try:
        zf = zipfile.ZipFile(BytesIO(zip_bytes))
        for fn in zf.namelist():
            if not fn.endswith('/'):
                ext = Path(fn).suffix.lower()
                if ext in ['.docx','.doc','.pdf','.xlsx','.xls','.xlsm','.csv','.txt','.xml','.html','.rtf']:
                    all_files[fn] = zf.read(fn)
        zf.close()
    except: return {}
    
    if not all_files: return {}
    
    logger.info(f"  {ICON_ZIP} {len(all_files)} fichiers à analyser")
    
    avis_text = ""
    dt_text = ""
    other_texts = []
    
    for fn, fb in all_files.items():
        txt = extract_text_from_any(fn, fb)
        if not txt.strip(): continue
        
        if is_avis_fr(fn, txt) and not avis_text:
            avis_text = f"=== AVIS: {fn} ===\n{txt[:8000]}"
            logger.info(f"  {ICON_AVIS} Avis: {fn} ({len(txt):,}c)")
        elif is_rc_file(fn):
            dt = extract_dossier_technique(txt)
            if dt and (not dt_text or len(dt) > len(dt_text)):
                dt_text = f"=== DOSSIER TECHNIQUE ({fn}) ===\n{dt}"
                logger.info(f"  {ICON_RC} Dossier Technique: {fn} ({len(dt):,}c)")
        else:
            if len(txt) > 200:
                other_texts.append(f"=== {fn} ===\n{txt[:1000]}")
    
    # Si pas de dossier technique trouvé, chercher dans tous les fichiers
    if not dt_text:
        for fn, fb in all_files.items():
            txt = extract_text_from_any(fn, fb)
            if not txt.strip(): continue
            dt = extract_dossier_technique(txt)
            if dt:
                dt_text = f"=== DOSSIER TECHNIQUE ({fn}) ===\n{dt}"
                logger.info(f"  {ICON_RC} Dossier Technique (fallback): {fn} ({len(dt):,}c)")
                break
    
    combined = avis_text
    if dt_text: combined += "\n\n" + dt_text
    if other_texts: combined += "\n\n" + "\n".join(other_texts[:3])
    
    if not combined.strip(): return {"error": "Aucun texte extrait"}
    
    total_chars = len(combined)
    logger.info(f"  {ICON_AI} Envoi Groq: {total_chars:,} caractères")
    
    result = extract_with_groq(combined, tender_ref)
    if result:
        result["_chars"] = total_chars
    return result


# ═══════════════ SINGLE ═══════════════

def analyze_single(ref):
    print(f"\n  Réf: {ref}")
    tenders = _sb_get_tenders({"select":"reference,objet,dce_zip_url","reference":f"eq.{ref}","limit":"1"})
    if not tenders: return print(f"  {ICON_KO} Non trouvée")
    t = tenders[0]
    if not t.get("dce_zip_url"): return print(f"  {ICON_KO} Pas de DCE")
    
    print(f"  Objet: {(t.get('objet') or '')[:120]}")
    zb = download_zip(t["dce_zip_url"])
    if not zb: return print(f"  {ICON_KO} Download échoué")
    
    print(f"  {ICON_ZIP} {len(zb):,} bytes")
    r = process_dce(zb, ref)
    
    if not r or "error" in r:
        return print(f"  {ICON_WARN} Rien extrait")
    
    print(f"\n  {'─'*40}\n  RÉSULTATS:\n  {'─'*40}")
    for field, icon, label, unit in [
        ("estimation","💰","Estimation","DH"),
        ("caution_provisoire","🔒","Caution provisoire","DH"),
        ("objet","📋","Objet",""),
        ("date_soumission","📅","Date",""),
        ("classe_demandee","📜","Classe",""),
        ("attestation_reference_demandee","📜","Références",""),
        ("visite_lieux_obligatoire","🏗️","Visite lieux",""),
        ("chiffre_affaires","💵","Chiffre d'affaires","DH"),
    ]:
        val = r.get(field)
        if val is not None and val != "":
            d = "OBLIGATOIRE" if val is True else "Facultative" if val is False else f"{val:,.0f} {unit}" if isinstance(val,(int,float)) else str(val)
            print(f"  {icon} {label}: {d}")
        else:
            print(f"  {icon} {label}: —")
    
    # 🔥 SAUVEGARDE CORRIGÉE — mapping vers tenders_3
    field_mapping = {
        "estimation": "estimation",
        "caution_provisoire": "caution_provisoire",
        "classe_demandee": "classe_demandee",
        "attestation_reference_demandee": "attestation_reference_demandee",
        "visite_lieux_obligatoire": "visite_lieux_obligatoire",
    }
    
    db_fields = {}
    for src, dst in field_mapping.items():
        val = r.get(src)
        if val is not None and val != "":
            db_fields[dst] = val
    
    if db_fields:
        if _sb_patch_tender(ref, db_fields):
            print(f"\n  {ICON_SAVE} {ICON_OK} Sauvegardé ({len(db_fields)} champs)")
        else:
            print(f"\n  {ICON_SAVE} {ICON_KO} Échec — vérifiez les colonnes")
    else:
        print(f"\n  {ICON_WARN} Rien à sauvegarder")


# ═══════════════ MAIN ═══════════════

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="DCE Extraction v13 — Fixed")
    parser.add_argument("--reference", "-r", type=str)
    args = parser.parse_args()
    
    print(f"\n{'='*60}")
    print(f"  {ICON_AI} CrystalWater — DCE Analyzer v13 (Fixed)")
    print(f"  Groq: {'✅ ' + GROQ_MODEL if GROQ_API_KEY else '❌ OFF'}")
    print(f"{'='*60}")
    
    if args.reference:
        analyze_single(args.reference)
    else:
        ref = input("\n  Référence: ").strip()
        if ref: analyze_single(ref)