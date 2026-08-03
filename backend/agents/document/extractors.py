"""
extractors.py — File extractors for PDF, Word, Excel
"""

import io
import logging
from typing import Dict, Any, List
from pathlib import Path

logger = logging.getLogger("extractors")


# ─── PDF EXTRACTOR ──────────────────────────────────────────────────────

def extract_pdf_text(file_bytes: bytes) -> Dict[str, Any]:
    """Extract full text content from PDF using PyMuPDF."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        all_text = ""
        pages = []
        for page_num, page in enumerate(doc):
            text = page.get_text()
            all_text += text + "\n\n"
            pages.append({
                "page_num": page_num + 1,
                "text": text,
                "word_count": len(page.get_text("words")),
                "has_images": len(page.get_images(full=True)) > 0
            })
        doc.close()
        return {
            "text": all_text,
            "pages": pages,
            "page_count": len(pages),
            "total_words": sum(p["word_count"] for p in pages),
            "type": "pdf_native"
        }
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return {"text": "", "pages": [], "page_count": 0, "total_words": 0, "type": "pdf_native", "error": str(e)}


# ─── OCR EXTRACTOR ─────────────────────────────────────────────────────

def extract_pdf_ocr(file_bytes: bytes) -> Dict[str, Any]:
    """Extract text from scanned PDF using OCR."""
    try:
        import pytesseract
        from PIL import Image
        import pdf2image

        images = pdf2image.convert_from_bytes(file_bytes, dpi=300)
        all_text = ""
        pages = []
        for page_num, img in enumerate(images):
            text = pytesseract.image_to_string(img, lang='fra+eng')
            all_text += text + "\n\n"
            pages.append({
                "page_num": page_num + 1,
                "text": text,
                "word_count": len(text.split())
            })
        return {
            "text": all_text,
            "pages": pages,
            "page_count": len(pages),
            "total_words": sum(p["word_count"] for p in pages),
            "type": "pdf_ocr"
        }
    except ImportError as e:
        logger.warning(f"OCR library missing: {e}")
        return {"text": "", "pages": [], "page_count": 0, "total_words": 0, "type": "pdf_ocr", "error": f"Missing library: {e}"}
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        return {"text": "", "pages": [], "page_count": 0, "total_words": 0, "type": "pdf_ocr", "error": str(e)}


# ─── WORD EXTRACTOR (FIXED: now reads tables too) ─────────────────────

def _iter_table_text(table) -> List[str]:
    """Recursively pull text out of a docx table, including nested tables."""
    lines = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            cell_text = cell.text.strip() if cell.text else ""
            if cell_text:
                row_cells.append(cell_text)
            # Handle nested tables inside a cell
            for nested_table in cell.tables:
                lines.extend(_iter_table_text(nested_table))
        if row_cells:
            lines.append(" | ".join(row_cells))
    return lines


def extract_word_text(file_bytes: bytes, ext: str = ".docx") -> Dict[str, Any]:
    """Extract text from Word document, including paragraphs AND tables."""
    try:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))

            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

            # Tables often hold the actual Avis/RC field values
            for table in doc.tables:
                text_parts.extend(_iter_table_text(table))

            full_text = "\n".join(text_parts)
            return {
                "text": full_text,
                "pages": [],
                "page_count": 0,
                "total_words": len(full_text.split()),
                "type": "word"
            }
        except ImportError:
            return {"text": "", "pages": [], "page_count": 0, "total_words": 0, "type": "word", "error": "python-docx not installed"}
    except Exception as e:
        return {"text": "", "pages": [], "page_count": 0, "total_words": 0, "type": "word", "error": str(e)}


# ─── EXCEL EXTRACTOR ───────────────────────────────────────────────────

def extract_excel_text(file_bytes: bytes, ext: str = ".xlsx") -> Dict[str, Any]:
    """Extract text from Excel document."""
    try:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
            text_parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                text_parts.append(f"--- Sheet: {sheet_name} ---")
                for row in ws.iter_rows(min_row=1, max_row=50, values_only=True):
                    row_text = " ".join(str(cell) for cell in row if cell is not None and str(cell).strip())
                    if row_text.strip():
                        text_parts.append(row_text)
            full_text = "\n".join(text_parts)
            return {
                "text": full_text,
                "pages": [],
                "page_count": 0,
                "total_words": len(full_text.split()),
                "type": "excel"
            }
        except ImportError:
            return {"text": "", "pages": [], "page_count": 0, "total_words": 0, "type": "excel", "error": "openpyxl not installed"}
    except Exception as e:
        return {"text": "", "pages": [], "page_count": 0, "total_words": 0, "type": "excel", "error": str(e)}


# ─── ROUTER ─────────────────────────────────────────────────────────────

def route_extractor(file_bytes: bytes, file_info: Dict[str, Any]) -> Dict[str, Any]:
    """Route file to the right extractor. Skips unknown file types."""
    file_type = file_info.get("type", "unknown")
    is_scanned = file_info.get("is_scanned", False)
    ext = file_info.get("ext", "")

    logger.info(f"Routing: {file_type} | scanned={is_scanned} | ext={ext}")

    if file_type not in ["pdf", "word", "excel"]:
        logger.info(f"Skipping unknown file type: {file_type}")
        return {
            "text": "",
            "pages": [],
            "page_count": 0,
            "total_words": 0,
            "type": file_type,
            "skipped": True,
            "skip_reason": f"Unsupported file type: {file_type}"
        }

    if file_type == "pdf":
        if is_scanned:
            result = extract_pdf_ocr(file_bytes)
        else:
            result = extract_pdf_text(file_bytes)
        result["skipped"] = False
        return result

    elif file_type == "word":
        result = extract_word_text(file_bytes, ext)
        result["skipped"] = False
        return result

    elif file_type == "excel":
        result = extract_excel_text(file_bytes, ext)
        result["skipped"] = False
        return result

    return {
        "text": "",
        "pages": [],
        "page_count": 0,
        "total_words": 0,
        "type": file_type,
        "skipped": True,
        "skip_reason": f"Unhandled file type: {file_type}"
    }