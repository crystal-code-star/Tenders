"""
00_rc_Extraction.py — RC Field Extraction — v47 FINAL
======================================================
CrystalWater — Traitement d'eau & Refroidissement industriel
https://crystalwater.ma/

Pipeline:
- Fichiers Word (.docx/.doc) + PDF natif → Regex
- PDF RC scanné (30-50 pages) → OCR toutes les pages → Groq LLM → Regex fallback
- Extraction de 18 champs + stockage Supabase automatique
- Mode batch : --all pour traiter toutes les références

Usage:
  python agents/document/00_rc_Extraction.py --reference "REFERENCE"
  python agents/document/00_rc_Extraction.py --all
  python agents/document/00_rc_Extraction.py --all --limit 10
"""

import os
import re
import sys
import json
import time
import logging
import argparse
import tempfile
import zipfile
import subprocess
from io import BytesIO
from pathlib import Path
from typing import Optional, Tuple
from datetime import datetime, timezone
from dotenv import load_dotenv
from colorama import init, Fore, Style
import supabase
import requests as http_requests

init(autoreset=True)

current_dir = Path(__file__).resolve().parent
sys.path.insert(0, str(current_dir))
env_path = current_dir.parent.parent / ".env"
load_dotenv(env_path)

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")
supabase_client = supabase.create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY) if SUPABASE_URL and SUPABASE_SERVICE_KEY else None

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")

logging.basicConfig(level=logging.WARNING, format="%(message)s")
logger = logging.getLogger(__name__)

# ─── Icons ─────────────────────────────────────────────────────
ICON_RC = "📜"
ICON_REF, ICON_TYPE_REF, ICON_NB_REF = "📎", "📋", "🔢"
ICON_CERTIF, ICON_CA = "🏷️", "💰"
ICON_DECL, ICON_CAUT, ICON_NOTE = "📝", "🔒", "👷"
ICON_PROSP, ICON_PLAN = "📦", "📊"
ICON_METHOD, ICON_MEMOIRE, ICON_ECHANT = "📐", "📑", "🧪"
ICON_ACTE, ICON_BORD = "📜", "📋"
ICON_ATT_CA, ICON_ATT_REF = "💰", "📎"
ICON_CHECK, ICON_CROSS, ICON_WARN = "✅", "❌", "⚠️"
ICON_ZIP, ICON_DOWNLOAD = "📦", "📥"
ICON_PDF, ICON_SCANNED, ICON_OCR, ICON_AI = "📄", "🖼️", "👁️", "🤖"
ICON_REGEX, ICON_DB = "🔍", "💾"
ICON_BATCH = "🔄"

# ─── Helpers ───────────────────────────────────────────────────
def print_info(msg): print(f"  📌 {Fore.WHITE}{msg}")
def print_success(msg): print(f"  {Fore.GREEN}✅ {msg}")
def print_error(msg): print(f"  {Fore.RED}❌ {msg}")
def print_warning(msg): print(f"  {Fore.YELLOW}⚠️  {msg}")


# ═══════════════════════════════════════════════════════════════
#  1. TEXT EXTRACTION (Word + PDF) — OPTIMISÉ LONGS DOCUMENTS
# ═══════════════════════════════════════════════════════════════

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells: parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception: return ""


def extract_text_from_doc(file_bytes: bytes) -> str:
    text = extract_text_from_docx(file_bytes)
    if len(text.strip()) > 200: return text
    try:
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(file_bytes); tmp_path = tmp.name
        output_dir = tempfile.mkdtemp()
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir', output_dir, tmp_path],
                       capture_output=True, text=True, timeout=60)
        for f in os.listdir(output_dir):
            if f.endswith('.txt'):
                with open(os.path.join(output_dir, f), 'r', encoding='utf-8', errors='ignore') as ff: text = ff.read().strip()
                break
        os.unlink(tmp_path)
        for f in os.listdir(output_dir):
            try: os.unlink(os.path.join(output_dir, f))
            except: pass
        os.rmdir(output_dir)
        if len(text.strip()) > 200: return text
    except: pass
    try:
        text = file_bytes.decode('latin-1', errors='ignore')
        text = re.sub(r'[^\x20-\x7E\xA0-\xFF\n\r\t\.\,\;\:\!\?\(\)\[\]\{\}\-\+\/\@\#\$\%\^\&\*\=]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
    except: pass
    return text


def ocr_page_with_tesseract(page, page_num: int = 0) -> str:
    """OCR d'une page avec prétraitement optimisé"""
    try:
        import pytesseract
        from PIL import Image, ImageFilter, ImageEnhance
        import numpy as np
        
        pix = page.get_pixmap(dpi=250)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img = img.convert('L')
        img = ImageEnhance.Contrast(img).enhance(2.5)
        img = img.filter(ImageFilter.SHARPEN)
        
        # Binarisation adaptative
        img_array = np.array(img)
        threshold = np.mean(img_array) * 0.8
        img_array = np.where(img_array > threshold, 255, 0).astype(np.uint8)
        img = Image.fromarray(img_array)
        
        config = '--oem 3 --psm 4 -c preserve_interword_spaces=1'
        text = pytesseract.image_to_string(img, lang='fra', config=config)
        return text
        
    except ImportError: return ""
    except Exception as e:
        logger.debug(f"OCR page {page_num}: {e}")
        return ""


def extract_text_from_pdf_rc(file_bytes: bytes, filename: str = "") -> Tuple[str, bool, str, int]:
    """
    Extraction optimisée pour longs PDF RC (30-50 pages)
    Returns: (text, is_scanned, diagnostic, total_pages)
    """
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = len(doc)
        
        print_info(f"PDF RC: {total_pages} pages")
        
        # Test rapide pages 1-3
        native_parts = []
        total_chars = 0
        test_pages = min(3, total_pages)
        
        for i in range(test_pages):
            t = doc[i].get_text("text").strip()
            if t:
                native_parts.append(t)
                total_chars += len(t)
        
        avg_chars = total_chars / max(test_pages, 1)
        is_scanned = total_chars < 100 or avg_chars < 50
        
        if not is_scanned:
            # PDF natif
            print_info("PDF natif - extraction texte...")
            for i in range(test_pages, total_pages):
                t = doc[i].get_text("text").strip()
                if t:
                    native_parts.append(t)
                    total_chars += len(t)
            doc.close()
            full_text = "\n".join(native_parts)
            print_success(f"Texte extrait: {len(full_text):,} caractères")
            return full_text, False, f"📄 PDF natif: {total_pages}p, {len(full_text):,} chars", total_pages
        
        # PDF scanné - OCR complet
        print_warning(f"PDF scanné - OCR {total_pages} pages...")
        ocr_parts = []
        total_ocr_chars = 0
        batch_size = 5
        
        for start_page in range(0, total_pages, batch_size):
            end_page = min(start_page + batch_size, total_pages)
            
            for i in range(start_page, end_page):
                page = doc[i]
                ocr_text = ocr_page_with_tesseract(page, i)
                
                if ocr_text.strip():
                    ocr_parts.append(f"--- Page {i+1} ---\n{ocr_text.strip()}")
                    total_ocr_chars += len(ocr_text)
                
                if total_pages > 10 and (i + 1) % 10 == 0:
                    print_info(f"OCR: {i+1}/{total_pages} pages ({total_ocr_chars:,} chars)")
            
            if end_page < total_pages:
                time.sleep(0.5)
        
        doc.close()
        ocr_text = "\n".join(ocr_parts)
        
        if ocr_text.strip():
            print_success(f"OCR terminé: {len(ocr_text):,} caractères")
            return ocr_text, True, f"🖼️ Scanné → OCR: {total_pages}p, {len(ocr_text):,} chars", total_pages
        else:
            return "", True, "🖼️ Scanné → OCR échoué", total_pages
            
    except ImportError:
        return "", False, "PyMuPDF manquant", 0
    except Exception as e:
        return "", False, f"Erreur PDF: {e}", 0


# ═══════════════════════════════════════════════════════════════
#  2. IDENTIFICATION FICHIERS RC
# ═══════════════════════════════════════════════════════════════

RC_NAME_PATTERNS = [
    r'\brcdp\b', r'\brcdg\b', r'\brc\b', r'^rc[\s_\-\.]', r'[\s_\-]rc[\s_\-\.]', r'[\s_\-]rc$',
    r'\br[èe]glement\s+de\s+consultation\b', r'\breglement\s+de\s+consultation\b',
    r'[_-]rcdp[_-]', r'[_-]rcdp$', r'^rcdp[_-]',
    r'[_-]rcdg[_-]', r'[_-]rcdg$', r'^rcdg[_-]',
    r'[_-]rc[_-]', r'[_-]rc\.', r'[_-]rc$',
]

def is_rc_file(filename: str) -> bool:
    name_lower = filename.lower(); name_no_ext = Path(filename).stem.lower()
    for p in RC_NAME_PATTERNS:
        if re.search(p, name_lower) or re.search(p, name_no_ext): return True
    return False

def is_supported_format(filename: str) -> bool:
    return filename.lower().endswith(('.docx', '.doc', '.pdf'))


# ═══════════════════════════════════════════════════════════════
#  3. DÉTECTION ACHETEUR
# ═══════════════════════════════════════════════════════════════

def detect_acheteur(text: str) -> str:
    if re.search(r'ONEE|Office\s+National\s+de\s+l[\u2019\']\u00e9lectricit', text, re.IGNORECASE): return "ONEE"
    if re.search(r'SRM|Soci[\u00e9e]t[\u00e9e]\s+R[\u00e9e]gionale\s+Multiservices', text, re.IGNORECASE): return "SRM"
    if re.search(r'ORMVA|Office\s+R[\u00e9e]gional\s+de\s+Mise\s+en\s+Valeur\s+Agricole', text, re.IGNORECASE): return "ORMVAO"
    if re.search(r'Conseil\s+(?:de\s+la\s+)?R[\u00e9e]gion|R[\u00e9e]gion\s+de', text, re.IGNORECASE): return "REGION"
    if re.search(r'Commune|Conseil\s+Communal|Pr[\u00e9e]sident\s+du\s+conseil', text, re.IGNORECASE): return "COMMUNE"
    return "AUTRE"


# ═══════════════════════════════════════════════════════════════
#  4. GROQ LLM — OPTIMISÉ LONGS DOCUMENTS
# ═══════════════════════════════════════════════════════════════

SYSTEM_PROMPT_RC = """Tu es un expert en analyse de RÈGLEMENTS DE CONSULTATION (RC) marocains.
Extrais ces 18 informations du texte. Réponds UNIQUEMENT avec ce JSON valide:

{
  "attestations_demandees": "OUI/NON",
  "types_attestations": "description ou null",
  "nombre_references": "nombre ou null",
  "classe_qualification": "ex: Classe A ou null",
  "chiffre_affaires": "ex: 30% du montant ou 500000 DH ou null",
  "declaration_honneur": "OUI/NON",
  "caution_provisoire": "OUI/NON",
  "note_moyens_humains": "OUI/NON",
  "depot_prospectus": "PHYSIQUE/ELECTRONIQUE/NON/null",
  "plan_charge": "OUI/NON",
  "moyens_humains_techniques": "OUI/NON",
  "methodologie_travail": "OUI/NON",
  "memoire_technique": "OUI/NON",
  "echantillon": "OUI/NON",
  "acte_engagement": "OUI/NON",
  "bordereau_prix": "OUI/NON",
  "attestations_ca": "OUI/NON",
  "attestations_reference": "OUI/NON"
}

RÈGLES STRICTES:
- Si l'information est présente: mets la valeur exacte
- Si l'information est ABSENTE: mets null (pas "NON", pas vide, null)
- Pour OUI/NON: "OUI" seulement si explicitement exigé
- Ne JAMAIS inventer
"""


def _extract_chunk_with_groq(text: str, filename: str, chunk_num: int, total_chunks: int) -> dict:
    """Extraction Groq pour un chunk"""
    if len(text.strip()) < 200:
        return {}
    
    try:
        response = http_requests.post(
            GROQ_API_URL,
            headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
            json={
                "model": GROQ_MODEL,
                "messages": [
                    {"role": "system", "content": SYSTEM_PROMPT_RC},
                    {"role": "user", "content": f"Fichier: {filename} (partie {chunk_num}/{total_chunks})\n\nTexte:\n{text[:10000]}"}
                ],
                "temperature": 0.1,
                "max_tokens": 800
            },
            timeout=90
        )
        
        if response.status_code == 200:
            content = response.json()["choices"][0]["message"]["content"].strip()
            content = content.replace("```json", "").replace("```", "").strip()
            
            json_match = re.search(r'\{[^{}]*\}', content, re.DOTALL)
            if json_match:
                try: return json.loads(json_match.group(0))
                except: pass
            try: return json.loads(content)
            except: pass
        
        return {}
    except Exception as e:
        logger.debug(f"Groq chunk {chunk_num} error: {e}")
        return {}


def _merge_chunk_results(all_results: list) -> dict:
    """Fusionne les résultats de plusieurs chunks"""
    if not all_results:
        return {}
    
    merged = {}
    
    boolean_fields = [
        'attestations_demandees', 'declaration_honneur', 'caution_provisoire',
        'note_moyens_humains', 'plan_charge', 'moyens_humains_techniques',
        'methodologie_travail', 'memoire_technique', 'echantillon',
        'acte_engagement', 'bordereau_prix', 'attestations_ca', 'attestations_reference'
    ]
    
    for field in boolean_fields:
        for result in all_results:
            try:
                val = str(result.get(field, '')).upper() if result.get(field) else ''
                if val == 'OUI':
                    merged[field] = 'OUI'
                    break
                elif val == 'NON':
                    if field not in merged:
                        merged[field] = 'NON'
            except Exception:
                continue
    
    text_fields = ['types_attestations', 'classe_qualification', 'chiffre_affaires', 'nombre_references', 'depot_prospectus']
    
    for field in text_fields:
        best_value = None
        best_length = 0
        for result in all_results:
            try:
                val = result.get(field)
                if val and str(val).strip() and str(val).lower() not in ['null', 'none', '']:
                    if len(str(val)) > best_length:
                        best_length = len(str(val))
                        best_value = str(val)
            except Exception:
                continue
        if best_value:
            merged[field] = best_value
    
    return merged


def extract_with_groq_rc(text: str, filename: str, total_pages: int = 0) -> dict:
    """Extraction Groq avec chunking pour longs documents"""
    if not GROQ_API_KEY or not text.strip():
        return {}
    
    print_info(f"🤖 Analyse Groq ({len(text):,} caractères)...")
    
    if len(text) <= 12000:
        result = _extract_chunk_with_groq(text, filename, 1, 1)
        if result:
            found = sum(1 for v in result.values() if v and str(v).lower() not in ['null', 'none', ''])
            print_success(f"Groq: {found}/18 champs extraits")
        return result if result else {}
    
    # Découpage intelligent
    print_info("Document long - découpage en sections...")
    pages = re.split(r'--- Page \d+ ---', text)
    
    if len(pages) <= 1:
        chunks = [text[i:i+8000] for i in range(0, len(text), 8000)]
    else:
        chunks = []
        current = ""
        for page in pages:
            if len(current) + len(page) > 8000:
                if current:
                    chunks.append(current)
                current = page
            else:
                current += "\n" + page
        if current:
            chunks.append(current)
    
    # Limiter à 8 chunks max
    chunks = chunks[:8]
    print_info(f"{len(chunks)} sections à analyser")
    
    all_results = []
    for i, chunk in enumerate(chunks):
        if len(chunk.strip()) < 300:
            continue
        result = _extract_chunk_with_groq(chunk, filename, i+1, len(chunks))
        if result:
            all_results.append(result)
        if len(chunks) > 2:
            time.sleep(0.5)
    
    merged = _merge_chunk_results(all_results)
    found = sum(1 for v in merged.values() if v and str(v).lower() not in ['null', 'none', ''])
    print_success(f"Groq fusionné: {found}/18 champs")
    return merged


# ═══════════════════════════════════════════════════════════════
#  5. EXTRACTION REGEX (18 CHAMPS) — CORRIGÉE
# ═══════════════════════════════════════════════════════════════

NUMBER_WORDS = {
    "un":1,"une":1,"deux":2,"trois":3,"quatre":4,"cinq":5,"six":6,"sept":7,"huit":8,"neuf":9,"dix":10,
    "onze":11,"douze":12,"treize":13,"quatorze":14,"quinze":15,
}

def parse_number_word(word: str) -> int:
    word = word.strip().lower()
    if word.isdigit(): return int(word)
    return NUMBER_WORDS.get(word, 0)


def extract_all_fields(text: str) -> dict:
    """Extrait les 18 champs depuis le texte."""
    result = {}
    normalized_flat = re.sub(r'\s+', ' ', text.replace('\r', ' '))

    section_finale = ""
    for marker_pat in [
        r'Pi\u00e8ce\s+1\s*:\s*R\u00e8glement\s+de\s+Consultation',
        r'Piece\s+1\s*:\s*Reglement\s+de\s+Consultation',
        r'Pi\u00e8ce\s+1\s*:\s*RC\b',
    ]:
        m = re.search(r'(' + marker_pat + r'\s*.*$)', normalized_flat, re.IGNORECASE)
        if m: section_finale = m.group(1).strip(); break
    if not section_finale:
        section_finale = normalized_flat[-3000:] if len(normalized_flat) > 3000 else normalized_flat

    # 1. ATTESTATIONS DE RÉFÉRENCE DEMANDÉES
    attestation_patterns = [
        r'attestations?\s+de\s+fin\s+d[\u2019\']ex[\u00e9e]cution',
        r'(?:Les\s+)?attestations?\s+ou\s+(?:leurs?|Ieurs?)\s+copies\s+certifi[\u00e9e]es\s+conformes?\s+[\u00e0a]\s+l[\u2019\']original\s+d[\u00e9e]livr[\u00e9e]es?\s+par\s+(?:les\s+)?ma[\u00eei]tres?\s+d[\u2019\']ouvrage',
        r'attestations?\s+ou\s+Ieurs\s+copies\s+certifi[\u00e9e]es\s+conformes?\s+[\u00e0a]\s+l[\u2019\']original\s+d[\u00e9e]livr[\u00e9e]es?\s+par\s+(?:les\s+)?ma[\u00eei]tres?\s+d[\u2019\']ouvrage',
        r'r[\u00e9e]f[\u00e9e]rences?\s+(?:de\s+)?projets?\s+similaires?',
        r'attestations?\s+de\s+r[\u00e9e]f[\u00e9e]rence',
        r'[Aa]ttestation\s+r[\u00e9e]f[\u00e9e]rence',
        r'attestations?\s+de\s+bonne\s+ex[\u00e9e]cution',
        r'attestations?\s+ou\s+leurs\s+copies\s+certifi[\u00e9e]es\s+conformes?',
        r'attestations?\s+d[\u00e9e]livr[\u00e9e]es?\s+par\s+(?:le|les)\s+ma[\u00eei]tre',
    ]
    result["attestations_demandees"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in attestation_patterns) else "NON"

    # 2. TYPES D'ATTESTATIONS — CORRIGÉ (try/except sur les groupes)
    types_found = False
    
    # Pattern 1
    m = re.search(
        r'(?:Les\s+)?attestations?\s+de\s+fin\s+d[\u2019\']ex[\u00e9e]cution\s+ou\s+leurs?\s+copies\s+certifi[\u00e9e]es\s+conformes?\s+[\u00e0a]\s+l[\u2019\']original\s+'
        r'd[\u00e9e]livr[\u00e9e]es?\s+par\s+(.*?)(?:\.\s+(?:Chaque|Ces|Les|En\s+cas|Dans\s+le\s+cas|Le\s+concurrent|$)|\n\n)',
        normalized_flat, re.IGNORECASE
    )
    if m:
        try:
            types_text = re.sub(r'\s+', ' ', m.group(1).strip())
            if len(types_text) > 20 and any(kw in types_text.lower() for kw in ['maître', 'maitre', 'ouvrage', 'homme', 'art']):
                result["types_attestations"] = "Attestations délivrées par " + types_text[:500]
                types_found = True
        except (IndexError, AttributeError):
            pass

    # Pattern 2
    if not types_found:
        m = re.search(
            r'(?:Les\s+)?attestations?\s+ou\s+(?:leurs?|Ieurs?)\s+copies\s+certifi[\u00e9e]es\s+conformes?\s+[\u00e0a]\s+l[\u2019\']original\s+'
            r'd[\u00e9e]livr[\u00e9e]es?\s+par\s+(.*?)(?:\.\s+(?:Chaque|Ces|Les|En\s+cas|Dans\s+le\s+cas|N\.?B\.?|$)|\n\n)',
            normalized_flat, re.IGNORECASE
        )
        if m:
            try:
                types_text = re.sub(r'\s+', ' ', m.group(1).strip())
                if len(types_text) > 20 and any(kw in types_text.lower() for kw in ['maître', 'maitre', 'ouvrage', 'homme', 'art', 'public', 'privé', 'titulaire']):
                    result["types_attestations"] = "Attestations délivrées par " + types_text[:500]
                    types_found = True
            except (IndexError, AttributeError):
                pass

    # Pattern 3
    if not types_found:
        m = re.search(
            r'(?:Les\s+)?attestations?\s+ou\s+Ieurs\s+copies\s+certifi[\u00e9e]es\s+conformes?\s+[\u00e0a]\s+l[\u2019\']original\s+'
            r'd[\u00e9e]livr[\u00e9e]es?\s+par\s+(.*?)(?:\.\s+(?:Chaque|Ces|Les|En\s+cas|Dans\s+le\s+cas|N\.?B\.?|$)|\n\n)',
            normalized_flat, re.IGNORECASE
        )
        if m:
            try:
                types_text = re.sub(r'\s+', ' ', m.group(1).strip())
                if len(types_text) > 20 and any(kw in types_text.lower() for kw in ['maître', 'maitre', 'ouvrage', 'homme', 'art', 'public', 'privé', 'titulaire']):
                    result["types_attestations"] = "Attestations délivrées par " + types_text[:500]
                    types_found = True
            except (IndexError, AttributeError):
                pass

    # Pattern 4 — CORRIGÉ : utilise groups()[-1] au lieu de group(2)
    if not types_found:
        m = re.search(
            r'(?:attestations?\s+(?:de\s+(?:fin\s+d[\u2019\']ex[\u00e9e]cution|r[\u00e9e]f[\u00e9e]rence)|ou\s+leurs?\s+copies)[^.]*?)'
            r'd[\u00e9e]livr[\u00e9e]es?\s+par\s+(.*?)(?:\.\s+(?:Chaque|Ces|Les|En\s+cas|$)|\n\n)',
            normalized_flat, re.IGNORECASE
        )
        if m:
            try:
                groups = m.groups()
                types_text = re.sub(r'\s+', ' ', groups[-1].strip())
                if len(types_text) > 20 and any(kw in types_text.lower() for kw in ['maître', 'maitre', 'ouvrage', 'homme', 'art', 'public', 'privé', 'titulaire']):
                    result["types_attestations"] = "Attestations délivrées par " + types_text[:500]
            except (IndexError, AttributeError):
                pass

    # 3. NOMBRE DE RÉFÉRENCES
    search_texts = [section_finale] if section_finale else []
    if normalized_flat not in search_texts: search_texts.append(normalized_flat)
    ref_patterns = [
        r'[Nn]ombre\s+de\s+r[\u00e9e]f[\u00e9e]rences?\s+exig[\u00e9e]es?\s+.*?:\s*(\w+)\s+r[\u00e9e]f[\u00e9e]rences?',
        r'(\w+)\s+r[\u00e9e]f[\u00e9e]rences?\s*,?\s*comme\s+suit',
        r'au\s+moins\s+(\w+)\s*(?:\((\d+)\))?\s*(?:attestation|projet|r[\u00e9e]f[\u00e9e]rence)',
        r'au\s+moins\s+(\w+)\s+(?:projet|r[\u00e9e]f[\u00e9e]rence)\s+similaire',
        r'au\s+moins\s+(\d+)\s+r[\u00e9e]f[\u00e9e]rences?',
        r'justifier\s+(?:de\s+)?(?:au\s+moins\s+)?(\d+)\s+r[\u00e9e]f[\u00e9e]rences?',
        r'disposer\s+d[\u2019\']au\s+moins\s+(\d+)\s+r[\u00e9e]f[\u00e9e]rences?',
        r'minimum\s+(?:de\s+)?(\d+)\s+r[\u00e9e]f[\u00e9e]rences?',
    ]
    found_nb = False
    for search_text in search_texts:
        for pat in ref_patterns:
            m = re.search(pat, search_text, re.IGNORECASE)
            if m:
                for g in m.groups():
                    if g and g.isdigit(): nb = int(g)
                    elif g:
                        nb = parse_number_word(g)
                        if nb <= 0: continue
                    else: continue
                    if 1 <= nb <= 50: result["nombre_references"] = str(nb); found_nb = True; break
                if found_nb: break
        if found_nb: break

    # 4. CLASSE / QUALIFICATION
    if "classe_qualification" not in result:
        m = re.search(
            r'[Ss]ecteur\s*\|\s*[Cc]lasse\s*\|\s*[Qq]ualifications?\s*\n?\s*(\w+)\s*\|\s*(\w+)\s*\|\s*(\w+)',
            text.replace('\r', ' '), re.IGNORECASE
        )
        if m:
            result["classe_qualification"] = f"Secteur {m.group(1).upper()} | Classe {m.group(2).upper()} | Qualification {m.group(3).upper()}"

    if "classe_qualification" not in result:
        for pat in [
            r'[Ss]ecteur\s+(\w+)\s*\|\s*[Cc]lasse\s+(\w+)\s*\|\s*[Qq]ualification\s+(\w+)',
            r'[Cc]ertificat\s+de\s+qualification\s+et\s+de\s+classification.*?[Cc]lasse\s*:?\s*(\w+)',
            r'[Cc]ertificat\s+de\s+qualification.*?[Cc]lasse\s*:?\s*(\w+)',
            r'[Qq]ualification\s+(?:exig[\u00e9e]e\s+)?.*?[Cc]lasse\s*:?\s*(\w+)',
        ]:
            m = re.search(pat, normalized_flat, re.IGNORECASE)
            if m:
                groups = m.groups()
                if len(groups) == 3:
                    result["classe_qualification"] = f"Secteur {groups[0].upper()} | Qualification {groups[1].upper()} | Classe {groups[2].upper()}"; break
                elif len(groups) == 1:
                    val = groups[0].upper()
                    if val and len(val) <= 3:
                        result["classe_qualification"] = f"Classe {val}"; break

    # 5. CHIFFRE D'AFFAIRES
    ca_search_texts = [section_finale] if section_finale else []
    if normalized_flat not in ca_search_texts: ca_search_texts.append(normalized_flat)
    found_ca = False
    for search_text in ca_search_texts:
        for pat in [
            r"chiffre\s+d[\u2019']affaires?\s+moyen\s+annuel.*?sup[\u00e9e]rieur\s+ou\s+[\u00e9e]gale?\s+[\u00e0a]\s+(\d+)\s*%",
            r"chiffre\s+d[\u2019']affaires?\s+.*?[\u00e9e]gale?\s+[\u00e0a]\s+(\d+)\s*%",
            r"chiffre\s+d[\u2019']affaires?\s+.*?au\s+moins\s+(\d+)\s*%",
            r"chiffre\s+d[\u2019']affaires?\s+.*?minimum\s+(?:de\s+)?(\d+)\s*%",
            r"sup[\u00e9e]rieur\s+ou\s+[\u00e9e]gale?\s+[\u00e0a]\s+(\d+)\s*%\s*(?:de\s+l[\u2019']estimation|du\s+montant)",
        ]:
            m = re.search(pat, search_text, re.IGNORECASE)
            if m: result["chiffre_affaires"] = f"{m.group(1)}% du montant de l'estimation"; found_ca = True; break
        if found_ca: break
    if not found_ca:
        for search_text in ca_search_texts:
            m = re.search(
                r"chiffre\s+d[\u2019']affaires?\s+.*?([\d]{1,3}(?:[\s\.]\d{3})*(?:,\d{2})?)\s*(?:DH|DHS|dhs|MAD)",
                search_text, re.IGNORECASE
            )
            if m: result["chiffre_affaires"] = f"{m.group(1)} DH"; break

    # 6. DÉCLARATION SUR L'HONNEUR
    declaration_patterns = [
        r'd[\u00e9e]claration\s+sur\s+l[\u2019\']honneur',
        r'mod[u\u00e8e]le\s+de\s+(?:la\s+)?d[\u00e9e]claration\s+sur\s+l[\u2019\']honneur',
        r'formulaire\s+de\s+d[\u00e9e]claration\s+sur\s+l[\u2019\']honneur',
        r'annexe.*?d[\u00e9e]claration\s+sur\s+l[\u2019\']honneur',
        r'pi[u\u00e8e]ce.*?d[\u00e9e]claration\s+sur\s+l[\u2019\']honneur',
    ]
    result["declaration_honneur"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in declaration_patterns) else "NON"

    # 7. CAUTION PROVISOIRE
    caution_patterns = [
        r'caution(?:nement)?\s+provisoire',
        r'r[\u00e9e]c[\u00e9e]piss[\u00e9e]\s+(?:du|electronique)\s+cautionnement',
        r'caution\s+personnelle\s+et\s+solidaire',
        r'attestation\s+de\s+(?:la\s+)?caution',
        r'constitution\s+(?:du|d[\u2019]un)\s+cautionnement',
        r'garantie\s+de\s+soumission',
        r'cautionnement\s+(?:est|sera|doit\s+[\u00eatre]tre)\s+(?:exig[\u00e9e]|constitu[\u00e9e]|demand[\u00e9e]|obligatoire)',
    ]
    result["caution_provisoire"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in caution_patterns) else "NON"

    # 8. NOTE MOYENS HUMAINS
    note_patterns = [
        r'note\s+indiquant\s+les\s+moyens\s+humains?\s+(?:et\s+techniques?)?',
        r'moyens\s+humains?\s+et\s+techniques?\s+du\s+concurrent',
        r'note\s+(?:sur\s+)?(?:les\s+)?moyens\s+humains?\s+(?:et\s+techniques?)?',
        r'description\s+des\s+moyens\s+humains?\s+(?:et\s+techniques?)?',
        r'pr[\u00e9e]sentation\s+des\s+moyens\s+humains?\s+(?:et\s+techniques?)?',
        r'Une\s+note\s+indiquant\s+les\s+moyens\s+humains',
        r'Une\s+note\s+indiquant.*?(?:moyens\s+humains|humains?\s+et\s+techniques)',
        r'note\s+indiquant.*?moyens\s+humains',
    ]
    result["note_moyens_humains"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in note_patterns) else "NON"

    # 9. DÉPÔT PROSPECTUS
    negation_patterns = [
        r"(?:d[\u00e9e]p[oô]t|remise|d[e\u00e9]p[o\u00f4]t).*?(?:n['\u2019]est\s+pas|ne\s+sera\s+pas).*?(?:exig[\u00e9e]|pr[\u00e9e]vu|demand[\u00e9e]|autoris[\u00e9e])",
        r"[Ii]l\s+n['\u2019]est\s+pas\s+pr[\u00e9e]vu\s+de\s+(?:remise|d[e\u00e9]p[o\u00f4]t)",
        r"(?:remise|d[\u00e9e]p[oô]t).*?n['\u2019]est\s+pas\s+(?:exig[\u00e9e]|pr[\u00e9e]vu|demand[\u00e9e]|obligatoire|requis)",
    ]
    if any(re.search(p, normalized_flat, re.IGNORECASE) for p in negation_patterns):
        result["depot_prospectus"] = "NON"
    elif re.search(
        r'(?:remise|d[\u00e9e]p[oô]t|fourniture)\s+(?:des?\s+)?[\u00e9e]chantillons.*?(?:est\s+(?:exig[\u00e9e]|demand[\u00e9e]|obligatoire|requis)|doit\s+[\u00eatre]tre\s+(?:fourni|remis|d[\u00e9e]pos[\u00e9e]|pr[\u00e9e]sent[\u00e9e])|sera\s+(?:exig[\u00e9e]|demand[\u00e9e]|requis))',
        normalized_flat, re.IGNORECASE
    ):
        result["depot_prospectus"] = "PHYSIQUE"
    elif re.search(
        r'(?:remise|d[\u00e9e]p[oô]t|fourniture)\s+(?:des?\s+)?[\u00e9e]chantillons.*?(?:voie\s+[\u00e9e]lectronique|portail|en\s+ligne|d[\u00e9e]mat[\u00e9e]rialis[\u00e9e]|[\u00e9e]lectroniquement)',
        normalized_flat, re.IGNORECASE
    ):
        result["depot_prospectus"] = "ELECTRONIQUE"

    # 10-18: PLAN DE CHARGE, MOYENS HUMAINS, MÉTHODOLOGIE, MÉMOIRE TECHNIQUE, ÉCHANTILLON, ACTE D'ENGAGEMENT, BORDEREAU, ATTESTATIONS CA, ATTESTATIONS RÉFÉRENCE
    plan_charge_patterns = [
        r'(?:d[\u00e9e]claration|mod[u\u00e8e]le|formulaire)\s+(?:du\s+)?plan\s+de\s+charge',
        r'plan\s+de\s+charge\s+(?:conforme\s+au\s+mod[u\u00e8e]le|fourni|joint|en\s+annexe)',
        r'fournir\s+(?:un|le)\s+plan\s+de\s+charge',
        r'joindre\s+(?:un|le)\s+plan\s+de\s+charge',
        r'pr[\u00e9e]senter\s+(?:un|le)\s+plan\s+de\s+charge',
        r'plan\s+de\s+charge\s+(?:d[\u00e9e]taill[\u00e9e]|pr[\u00e9e]visionnel)',
        r'annexe.*?plan\s+de\s+charge',
        r'produire\s+(?:un|le)\s+plan\s+de\s+charge',
        r'remettre\s+(?:un|le)\s+plan\s+de\s+charge',
    ]
    result["plan_charge"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in plan_charge_patterns) else "NON"

    moyens_ht_patterns = [
        r'(?:liste|moyens|description)\s+(?:des\s+)?moyens\s+humains?\s+(?:et\s+techniques?|techniques?)?\s+(?:[àa]\s+affecter|mis\s+[àa]\s+(?:la\s+)?disposition|propos[\u00e9e]s|du\s+projet)',
        r'moyens\s+humains?\s+(?:et\s+techniques?|techniques?)?\s+(?:[àa]\s+affecter|mis\s+[àa]\s+(?:la\s+)?disposition|propos[\u00e9e]s)',
        r'moyens\s+mat[\u00e9e]riels?\s+(?:[àa]\s+affecter|mis\s+[àa]\s+(?:la\s+)?disposition|propos[\u00e9e]s|minimums?\s+(?:exig[\u00e9e]s?|requis|demand[\u00e9e]s?))',
        r'liste\s+(?:des\s+)?moyens\s+(?:humains|mat[\u00e9e]riels|humains?\s+et\s+mat[\u00e9e]riels)',
        r'moyens\s+humains?\s+(?:et\s+techniques?|techniques?)?\s+du\s+concurrent',
        r'(?:[\u00e9e]quipes?\s+de\s+(?:pose|g[\u00e9e]nie\s+civil)|nombre\s+d[\u2019\']?[\u00e9e]quipes?)',
        r'liste\s+(?:du\s+)?personnel\s+(?:cl[\u00e9e]|affect[\u00e9e]|propos[\u00e9e])',
        r'(?:curriculum\s+vitae|CV|C\.V\.)\s+(?:du\s+)?(?:directeur|chef|responsable|technicien)',
    ]
    result["moyens_humains_techniques"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in moyens_ht_patterns) else "NON"

    methodo_patterns = [
        r'm[\u00e9e]thodologie\s+(?:de\s+(?:travail|r[\u00e9e]alisation|mise\s+en\s+[\u0153u]vre)|propos[\u00e9e]e|d[\u00e9e]taill[\u00e9e]e|et\s+planning)',
        r'm[\u00e9e]thode\s+de\s+travail',
        r'approche\s+m[\u00e9e]thodologique',
        r'planning\s+(?:pr[\u00e9e]visionnel|de\s+r[\u00e9e]alisation|d[\u00e9e]taill[\u00e9e]|des\s+travaux)',
        r'note\s+m[\u00e9e]thodologique',
        r'description\s+de\s+la\s+m[\u00e9e]thodologie',
        r'm[\u00e9e]thodologie\s+et\s+planning\s+de\s+r[\u00e9e]alisation',
        r'calendrier\s+d[\u2019\']ex[\u00e9e]cution',
        r'chronogramme',
        r'(?:diagramme\s+de\s+)?[Gg]antt',
    ]
    result["methodologie_travail"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in methodo_patterns) else "NON"

    memoire_patterns = [
        r'm[\u00e9e]moire\s+technique',
        r'rapport\s+technique',
        r'document\s+technique\s+(?:d[\u00e9e]taill[\u00e9e]|explicatif|descriptif)',
        r'offre\s+technique\s+(?:est\s+exig[\u00e9e]e|doit\s+comprendre|comprend|comporte|comportera|sera\s+(?:exig[\u00e9e]e|demand[\u00e9e]e|requise))',
        r'dossier\s+technique\s+(?:d[\u00e9e]taill[\u00e9e]|complet|explicatif)',
        r'note\s+technique\s+(?:d[\u00e9e]taill[\u00e9e]e|descriptive|explicative)',
        r'sp[\u00e9e]cifications?\s+techniques?\s+(?:d[\u00e9e]taill[\u00e9e]es?|des\s+[u00e9e]quipements?|du\s+mat[\u00e9e]riel)',
        r'descriptif\s+technique',
        r'catalogue\s+technique',
    ]
    result["memoire_technique"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in memoire_patterns) else "NON"

    echantillon_patterns = [
        r'[\u00e9e]chantillons?\s+(?:est\s+(?:exig[\u00e9e]|demand[\u00e9e]|obligatoire|requis|n[\u00e9e]cessaire)|doit\s+[\u00eatre]tre\s+(?:fourni|remis|pr[\u00e9e]sent[\u00e9e]|d[\u00e9e]pos[\u00e9e])|sera\s+(?:exig[\u00e9e]|demand[\u00e9e]|requis))',
        r'(?:remise|d[\u00e9e]p[oô]t|fourniture)\s+(?:des?\s+)?[\u00e9e]chantillons?\s+(?:est|sera)\s+(?:exig[\u00e9e]|obligatoire|demand[\u00e9e]|requis|n[\u00e9e]cessaire)',
        r'prototypes?\s+(?:est\s+(?:exig[\u00e9e]|demand[\u00e9e]|obligatoire|requis)|doit\s+[\u00eatre]tre\s+(?:fourni|remis|pr[\u00e9e]sent[\u00e9e]|d[\u00e9e]pos[\u00e9e])|sera\s+(?:exig[\u00e9e]|demand[\u00e9e]|requis))',
        r'(?:[\u00e9e]chantillon|prototype|sp[\u00e9e]cimen)\s+(?:doit\s+[\u00eatre]tre|sera)\s+(?:fourni|remis|pr[\u00e9e]sent[\u00e9e]|d[\u00e9e]pos[\u00e9e]|envoy[\u00e9e]|transmis)',
        r'pr[\u00e9e]senter\s+(?:un|des|le|les)\s+(?:[\u00e9e]chantillon|prototype|sp[\u00e9e]cimen)',
        r'fournir\s+(?:un|des|le|les)\s+(?:[\u00e9e]chantillon|prototype|sp[\u00e9e]cimen)',
        r'joindre\s+(?:un|des|le|les)\s+(?:[\u00e9e]chantillon|prototype|sp[\u00e9e]cimen)',
        r'd[\u00e9e]p[oô]t\s+(?:des?\s+)?[\u00e9e]chantillons?\s+(?:est|sera|doit\s+[\u00eatre]tre)\s+(?:exig[\u00e9e]|demand[\u00e9e]|obligatoire|requis|effectu[\u00e9e]|r[\u00e9e]alis[\u00e9e])',
    ]
    result["echantillon"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in echantillon_patterns) else "NON"

    acte_patterns = [
        r'acte\s+d[\u2019\']engagement',
        r'mod[u\u00e8e]le\s+(?:de\s+l[\u2019\']|d[\u2019\'])?acte\s+d[\u2019\']engagement',
        r'l[\u2019\']acte\s+d[\u2019\']engagement\s+(?:par\s+lequel|du|est|doit|sera|établi|signé)',
        r'acte\s+d[\u2019\']engagement\s+d[ûu]ment\s+rempli',
        r'acte\s+d[\u2019\']engagement.*?mod[u\u00e8e]le\s+joint',
        r'formulaire\s+(?:de\s+l[\u2019\']|d[\u2019\'])?acte\s+d[\u2019\']engagement',
        r'acte\s+d[\u2019\']engagement.*?selon\s+le\s+mod[u\u00e8e]le',
        r'acte\s+d[\u2019\']engagement.*?(?:en\s+un\s+(?:seul\s+)?exemplaire|établi\s+en)',
        r'annexe.*?acte\s+d[\u2019\']engagement',
        r'\b[Aa]cte\s+d[\u2019\'][Ee]ngagement\b',
        r'acte\s+d[\u2019\']engagement.*?conforme\s+au\s+mod[u\u00e8e]le',
        r'acte\s+d[\u2019\']engagement.*?(?:joint|ci-joint|en\s+annexe)',
    ]
    result["acte_engagement"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in acte_patterns) else "NON"

    bordereau_patterns = [
        r'bordereau\s+(?:des?\s+)?prix',
        r'bordereau\s+(?:des?\s+)?prix\s*[-–—]\s*d[\u00e9e]tail\s+estimatif',
        r'bordereau\s+(?:des?\s+)?prix[-\s]?d[\u00e9e]tail\s+estimatif',
        r'mod[u\u00e8e]le\s+(?:du\s+)?bordereau\s+(?:des?\s+)?prix',
        r'bordereau\s+(?:des?\s+)?prix\s+et\s+(?:le\s+)?d[\u00e9e]tail\s+estimatif',
        r'bordereau\s+(?:des?\s+)?prix.*?(?:d[ûu]ment\s+rempli|rempli|compl[ée]t[ée]|renseign[ée])',
        r'bordereau\s+(?:des?\s+)?prix.*?sign[ée]',
        r'bordereau\s+(?:des?\s+)?prix.*?selon\s+le\s+mod[u\u00e8e]le',
        r'bordereau\s+(?:des?\s+)?prix.*?conforme\s+au\s+mod[u\u00e8e]le',
        r'bordereau\s+de\s+prix',
        r'\bBPU\b|\bB\.?P\.?\s*(?:des?\s+)?prix',
        r'd[\u00e9e]tail\s+estimatif.*?bordereau',
        r'bordereau\s+(?:des?\s+)?prix.*?article\s+30',
        r'bordereau\s+(?:des?\s+)?prix\s+unitaires?',
        r'bordereau.*?(?:mod[u\u00e8e]le\s+joint|en\s+annexe)',
        r'bordereau\s+(?:des?\s+)?prix.*?libell[ée]s?\s+en\s+chiffres',
        r'(?:les\s+)?prix\s+unitaires?\s+(?:du\s+)?bordereau',
        r'bordereau.*?d[\u00e9e]composition\s+du\s+montant\s+global',
    ]
    result["bordereau_prix"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in bordereau_patterns) else "NON"

    ca_attestation_patterns = [
        r'attestations?\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?',
        r'attestation\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?\s+r[\u00e9e]alis[\u00e9e]',
        r'attestations?\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?.*?d[\u00e9e]livr[\u00e9e]es?\s+par\s+(?:le\s+)?(?:Minist[u\u00e8e]re|DGI|Direction\s+G[\u00e9e]n[\u00e9e]rale\s+des\s+Imp[o\u00f4]ts)',
        r'attestations?\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?.*?(?:ann[\u00e9e]es?\s+\d{4}|exercices?\s+\d{4}|derni[u\u00e8e]res?\s+ann[\u00e9e]es)',
        r'chiffre\s+d[\u2019\']affaires?.*?attestation',
        r'attestation\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?.*?certifi[\u00e9e]e',
        r'attestation\s+de\s+chiffre\s+d[\u2019\']affaires?\s+annuel',
        r'attestations?\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?\s+ou\s+(?:leurs?|Ieurs?)\s+copies',
        r'capacit[\u00e9e]s?\s+financi[u\u00e8e]res?.*?attestation\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?',
        r'dossier\s+additif.*?attestations?\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?',
        r'attestation\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?.*?(?:trois|3)\s+(?:\(03\)\s+)?derni[u\u00e8e]res?\s+ann[\u00e9e]es',
        r'pi[u\u00e8e]ces?\s+justifiant\s+les\s+capacit[\u00e9e]s\s+financi[u\u00e8e]res?.*?attestation\s+(?:du\s+)?chiffre\s+d[\u2019\']affaires?',
    ]
    result["attestations_ca"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in ca_attestation_patterns) else "NON"

    ref_attestation_patterns = [
        r'attestations?\s+de\s+fin\s+d[\u2019\']ex[\u00e9e]cution',
        r'attestations?\s+de\s+r[\u00e9e]f[\u00e9e]rence',
        r'attestations?\s+de\s+bonne\s+ex[\u00e9e]cution',
        r'attestations?\s+ou\s+(?:leurs?|Ieurs?)\s+copies\s+certifi[\u00e9e]es\s+conformes?.*?(?:ma[\u00eei]tres?\s+d[\u2019\']ouvrage|hommes?\s+de\s+l[\u2019\']art)',
        r'(?:Les\s+)?attestations?.*?examin[\u00e9e]es?\s+par\s+la\s+commission',
        r'attestations?.*?crit[u\u00e8e]res?\s+d[\u2019\']admissibilit[\u00e9e]',
        r'attestations?.*?d[\u00e9e]livr[\u00e9e]es?\s+par\s+(?:les\s+)?ma[\u00eei]tres?\s+d[\u2019\']ouvrage\s*(?:publics?\s+ou\s+priv[\u00e9e]s?)',
        r'r[\u00e9e]f[\u00e9e]rences?\s+(?:de\s+)?projets?\s+similaires?',
        r'attestations?\s+[\u00e9e]tablies?\s+au\s+nom\s+d[\u2019\']un\s+groupement',
        r'[Cc]haque\s+attestation\s+pr[\u00e9e]cise\s+notamment',
        r'attestations?\s+de\s+fin\s+d[\u2019\']ex[\u00e9e]cution.*?num[\u00e9e]rot[\u00e9e]es?',
        r'attestations?.*?(?:ma[\u00eei]tre|maitre)\s+d[\u2019\']ouvrage',
        r'attestations?\s+(?:de\s+r[\u00e9e]f[\u00e9e]rence|de\s+fin\s+d[\u2019\']ex[\u00e9e]cution).*?(?:examin[\u00e9e]es?|v[\u00e9e]rifi[\u00e9e]es?|not[\u00e9e]es?)',
        r'dossier\s+technique.*?attestations?\s+(?:de\s+fin\s+d[\u2019\']ex[\u00e9e]cution|de\s+r[\u00e9e]f[\u00e9e]rence)',
    ]
    result["attestations_reference"] = "OUI" if any(re.search(p, normalized_flat, re.IGNORECASE) for p in ref_attestation_patterns) else "NON"

    return result


# ═══════════════════════════════════════════════════════════════
#  6. TRAITEMENT FICHIER RC
# ═══════════════════════════════════════════════════════════════

def process_rc_file(filename: str, file_bytes: bytes) -> dict:
    """Traite un fichier RC avec OCR optimisé pour longs documents"""
    ext = Path(filename).suffix.lower()
    if ext not in ('.docx', '.doc', '.pdf'):
        return {"filename": filename, "error": f"Format non supporté: {ext}"}
    
    is_scanned = False
    extraction_diag = ""
    text = ""
    total_pages = 0
    
    if ext == '.pdf':
        text, is_scanned, extraction_diag, total_pages = extract_text_from_pdf_rc(file_bytes, filename)
    elif ext == '.docx':
        text = extract_text_from_docx(file_bytes)
        extraction_diag = f"📝 DOCX: {len(text):,} chars"
    elif ext == '.doc':
        text = extract_text_from_doc(file_bytes)
        extraction_diag = f"📝 DOC: {len(text):,} chars"
    
    if is_scanned and not text.strip():
        return {
            "filename": filename,
            "error": "OCR échoué",
            "is_scanned": True,
            "extraction_diag": extraction_diag,
            "total_pages": total_pages
        }
    
    if not text or not text.strip():
        return {
            "filename": filename,
            "error": "Aucun texte extrait",
            "extraction_diag": extraction_diag
        }
    
    acheteur = detect_acheteur(text)
    method = ""
    extracted = {}
    
    all_keys = [
        'attestations_demandees', 'types_attestations', 'nombre_references',
        'classe_qualification', 'chiffre_affaires', 'declaration_honneur',
        'caution_provisoire', 'note_moyens_humains', 'depot_prospectus',
        'plan_charge', 'moyens_humains_techniques', 'methodologie_travail',
        'memoire_technique', 'echantillon', 'acte_engagement',
        'bordereau_prix', 'attestations_ca', 'attestations_reference'
    ]
    
    # Stratégie: OCR → Groq d'abord, puis Regex en fallback
    if is_scanned and GROQ_API_KEY:
        print_info("PDF scanné → Extraction Groq...")
        groq_result = extract_with_groq_rc(text, filename, total_pages)
        
        if groq_result:
            for key in all_keys:
                val = groq_result.get(key)
                if val and str(val).strip() and str(val).lower() not in ['null', 'none', '']:
                    extracted[key] = str(val).strip()
        
        if extracted:
            method = f"🤖 Groq LLM (OCR {total_pages}p)"
            # Compléter avec Regex pour les champs manquants
            try:
                regex_result = extract_all_fields(text)
                for key in all_keys:
                    if key not in extracted and regex_result.get(key):
                        extracted[key] = regex_result[key]
            except Exception as e:
                logger.debug(f"Regex fallback error: {e}")
        else:
            try:
                extracted = extract_all_fields(text)
                method = f"🔍 Regex (fallback OCR)"
            except Exception as e:
                logger.debug(f"Regex error: {e}")
                method = f"⚠️ Extraction échouée"
    else:
        try:
            extracted = extract_all_fields(text)
            method = f"🔍 Regex ({acheteur})"
        except Exception as e:
            logger.debug(f"Regex error: {e}")
            method = f"⚠️ Regex échoué"
    
    return {
        "filename": filename,
        "text_length": len(text),
        "extraction_diag": extraction_diag,
        "acheteur": acheteur,
        "method": method,
        "is_scanned": is_scanned,
        "total_pages": total_pages,
        **extracted
    }


def display_result(result: dict, index: int, total: int):
    fn = result.get('filename', '?'); acheteur = result.get('acheteur', '?')
    extraction_diag = result.get('extraction_diag', ''); method = result.get('method', '')
    if "error" in result: print(f"\n  {ICON_CROSS} {Fore.RED}[{index}/{total}] {fn}\n     {Fore.RED}{result['error']}"); return
    tl = result.get('text_length', 0); is_scanned = result.get('is_scanned', False)
    icon_file = ICON_SCANNED if is_scanned else ICON_RC
    print(f"\n  {Fore.CYAN}{'─'*70}")
    print(f"  {icon_file} {Fore.WHITE}{Style.BRIGHT}[{index}/{total}] {fn} ({tl:,} caractères)")
    if extraction_diag: print(f"     {Fore.LIGHTBLACK_EX}{extraction_diag}")
    print(f"     {Fore.LIGHTBLACK_EX}Acheteur: {acheteur} | Méthode: {method}")
    print(f"  {Fore.CYAN}{'─'*70}")

    val = result.get('attestations_demandees', '')
    print(f"  {ICON_REF} {Fore.GREEN}Attestations de référence demandées: {Fore.WHITE}{val} {'✅' if val == 'OUI' else '❌'}" if val else f"  {ICON_REF} {Fore.RED}Attestations de référence demandées: NON DÉTECTÉ")
    val = result.get('types_attestations', '')
    print(f"  {ICON_TYPE_REF} {Fore.GREEN}Types d'attestations: {Fore.WHITE}{str(val)[:200]}" if val else f"  {ICON_TYPE_REF} {Fore.RED}Types d'attestations: NON TROUVÉ")
    val = result.get('nombre_references', '')
    print(f"  {ICON_NB_REF} {Fore.GREEN}Nombre de références exigé: {Fore.WHITE}{val}" if val else f"  {ICON_NB_REF} {Fore.RED}Nombre de références exigé: NON TROUVÉ")
    for icon, label, field in [
        (ICON_CERTIF, "Certificat de qualification/classification", "classe_qualification"),
        (ICON_CA, "Chiffre d'affaires minimum exigé", "chiffre_affaires"),
    ]:
        val = result.get(field, '')
        print(f"  {icon} {Fore.GREEN}{label}: {Fore.WHITE}{str(val)[:300]}" if val else f"  {icon} {Fore.RED}{label}: NON TROUVÉ")
    for icon, label, field in [
        (ICON_DECL, "Déclaration sur l'honneur exigée", "declaration_honneur"),
        (ICON_CAUT, "Caution provisoire exigée", "caution_provisoire"),
        (ICON_NOTE, "Note moyens humains exigée", "note_moyens_humains"),
    ]:
        val = result.get(field, '')
        if val and str(val).upper() == 'OUI': print(f"  {icon} {Fore.GREEN}{label}: {Fore.WHITE}OUI ✅")
        elif val and str(val).upper() == 'NON': print(f"  {icon} {Fore.RED}{label}: {Fore.WHITE}NON ❌")
        else: print(f"  {icon} {Fore.YELLOW}{label}: NON DÉTECTÉ")
    val = result.get('depot_prospectus', '')
    if val == 'PHYSIQUE': print(f"  {ICON_PROSP} {Fore.YELLOW}Dépôt prospectus/notice/doc technique: {Fore.WHITE}PHYSIQUE 📦")
    elif val == 'ELECTRONIQUE': print(f"  {ICON_PROSP} {Fore.CYAN}Dépôt prospectus/notice/doc technique: {Fore.WHITE}ELECTRONIQUE 💻")
    elif val == 'NON': print(f"  {ICON_PROSP} {Fore.GREEN}Dépôt prospectus/notice/doc technique: {Fore.WHITE}NON ❌")
    else: print(f"  {ICON_PROSP} {Fore.YELLOW}Dépôt prospectus/notice/doc technique: NON DÉTECTÉ")
    val = result.get('plan_charge', '')
    if val and str(val).upper() == 'OUI': print(f"  {ICON_PLAN} {Fore.GREEN}Plan de charge exigé: {Fore.WHITE}OUI ✅")
    elif val and str(val).upper() == 'NON': print(f"  {ICON_PLAN} {Fore.RED}Plan de charge exigé: {Fore.WHITE}NON ❌")
    else: print(f"  {ICON_PLAN} {Fore.YELLOW}Plan de charge exigé: NON DÉTECTÉ")
    for icon, label, field in [
        (ICON_NOTE, "Moyens humains et techniques (liste)", "moyens_humains_techniques"),
        (ICON_METHOD, "Méthodologie de travail", "methodologie_travail"),
        (ICON_MEMOIRE, "Mémoire technique", "memoire_technique"),
        (ICON_ECHANT, "Échantillon/prototype", "echantillon"),
    ]:
        val = result.get(field, '')
        if val and str(val).upper() == 'OUI': print(f"  {icon} {Fore.GREEN}{label}: {Fore.WHITE}OUI ✅")
        elif val and str(val).upper() == 'NON': print(f"  {icon} {Fore.RED}{label}: {Fore.WHITE}NON ❌")
        else: print(f"  {icon} {Fore.YELLOW}{label}: NON DÉTECTÉ")
    for icon, label, field in [
        (ICON_ACTE, "Acte d'engagement exigé", "acte_engagement"),
        (ICON_BORD, "Bordereau des prix exigé", "bordereau_prix"),
    ]:
        val = result.get(field, '')
        if val and str(val).upper() == 'OUI': print(f"  {icon} {Fore.GREEN}{label}: {Fore.WHITE}OUI ✅")
        elif val and str(val).upper() == 'NON': print(f"  {icon} {Fore.RED}{label}: {Fore.WHITE}NON ❌")
        else: print(f"  {icon} {Fore.YELLOW}{label}: NON DÉTECTÉ")
    for icon, label, field in [
        (ICON_ATT_CA, "Attestations du chiffre d'affaires exigées", "attestations_ca"),
        (ICON_ATT_REF, "Attestations de référence (dossier) exigées", "attestations_reference"),
    ]:
        val = result.get(field, '')
        if val and str(val).upper() == 'OUI': print(f"  {icon} {Fore.GREEN}{label}: {Fore.WHITE}OUI ✅")
        elif val and str(val).upper() == 'NON': print(f"  {icon} {Fore.RED}{label}: {Fore.WHITE}NON ❌")
        else: print(f"  {icon} {Fore.YELLOW}{label}: NON DÉTECTÉ")


# ═══════════════════════════════════════════════════════════════
#  7. SUPABASE — RÉCUPÉRATION + STOCKAGE
# ═══════════════════════════════════════════════════════════════

def get_all_references(limit: int = None) -> list:
    if not supabase_client: print(f"{Fore.RED}❌ Supabase client not initialized"); return []
    try:
        query = supabase_client.table("tenders_3").select("reference").not_.is_("dce_zip_url", "null")
        if limit: query = query.limit(limit)
        resp = query.execute()
        return [r['reference'] for r in resp.data if r.get('reference')]
    except Exception as e: print(f"{Fore.RED}❌ Erreur: {e}"); return []


def get_rc_files_from_supabase(reference: str) -> list:
    if not supabase_client: return None
    try:
        resp = supabase_client.table("tenders_3").select("*").eq("reference", reference).execute()
        if not resp.data: print(f"\n{Fore.RED}❌ Référence non trouvée: {reference}"); return None
        t = resp.data[0]
        print(f"\n{Fore.GREEN}✅ {t.get('reference')} | {(t.get('objet') or '')[:120]}")
        zip_url = t.get('dce_zip_url'); base64_zip = t.get('dce_zip_base64'); zip_bytes = None
        if base64_zip: import base64; zip_bytes = base64.b64decode(base64_zip)
        elif zip_url:
            print(f"  {ICON_DOWNLOAD} Téléchargement du ZIP...")
            r = http_requests.get(zip_url, timeout=120)
            if r.status_code == 200: zip_bytes = r.content
            else: print(f"{Fore.RED}❌ Erreur HTTP {r.status_code}"); return None
        else: print(f"{Fore.RED}❌ Aucun ZIP disponible"); return None
        rc_files = []; ignored_not_supported = []; ignored_not_rc = []
        with zipfile.ZipFile(BytesIO(zip_bytes)) as zf:
            all_files = [fn for fn in zf.namelist() if not fn.endswith('/')]
            print(f"\n  {ICON_ZIP} {len(all_files)} fichiers dans le ZIP")
            print(f"  {Fore.CYAN}Filtrage : UNIQUEMENT RC/RCDP/RCDG en .docx/.doc/.pdf...")
            for fn in all_files:
                if not is_supported_format(fn): ignored_not_supported.append(fn); continue
                if not is_rc_file(fn): ignored_not_rc.append(fn); continue
                fb = zf.read(fn)
                icon = "📄" if fn.lower().endswith('.pdf') else "📝"
                rc_files.append({"filename": Path(fn).name, "file_bytes": fb, "size_kb": len(fb) / 1024})
                print(f"    {icon} {Fore.GREEN}{fn} ({len(fb)/1024:.1f} KB) ✅ RC")
        print(f"\n  {Fore.CYAN}{'─'*50}")
        print(f"  📊 Résumé du filtrage :")
        print(f"     {Fore.GREEN}✅ RC (Word/PDF) : {len(rc_files)}")
        print(f"     {Fore.YELLOW}⚠️  Word/PDF mais pas RC : {len(ignored_not_rc)}")
        print(f"     {Fore.LIGHTBLACK_EX}📁 Autres formats : {len(ignored_not_supported)}")
        if ignored_not_rc:
            print(f"\n  {Fore.YELLOW}Fichiers Word/PDF ignorés (pas RC) :")
            for fn in ignored_not_rc: print(f"     {Fore.LIGHTBLACK_EX}  {fn}")
        if not rc_files: print(f"\n{Fore.YELLOW}⚠️ Aucun fichier RC trouvé dans le ZIP"); return None
        print(f"\n{Fore.GREEN}✅ {len(rc_files)} fichier(s) RC à traiter")
        return rc_files
    except Exception as e:
        print(f"{Fore.RED}❌ Erreur: {e}")
        import traceback; traceback.print_exc()
        return None


def save_rc_extraction_to_supabase(reference: str, results: list):
    if not supabase_client: return False
    merged = {}; any_scanned = False
    for r in results:
        if "error" in r: continue
        if r.get('is_scanned'): any_scanned = True
        for key in ['attestations_demandees', 'types_attestations', 'nombre_references',
                   'classe_qualification', 'chiffre_affaires', 'declaration_honneur',
                   'caution_provisoire', 'note_moyens_humains', 'acheteur',
                   'depot_prospectus', 'plan_charge', 'moyens_humains_techniques',
                   'methodologie_travail', 'memoire_technique', 'echantillon',
                   'acte_engagement', 'bordereau_prix',
                   'attestations_ca', 'attestations_reference']:
            val = r.get(key, '')
            if val and not merged.get(key): merged[key] = val
    if not merged: return False
    def to_bool(val): return True if str(val).upper() == 'OUI' else False
    def to_int(val):
        try: return int(val)
        except: return None
    data = {
        "reference": reference,
        "attestations_demandees": to_bool(merged.get('attestations_demandees', 'NON')),
        "types_attestations": (merged.get('types_attestations') or '')[:2000],
        "nombre_references": to_int(merged.get('nombre_references', '')),
        "classe_qualification": (merged.get('classe_qualification') or '')[:1000],
        "chiffre_affaires": (merged.get('chiffre_affaires') or '')[:500],
        "declaration_honneur": to_bool(merged.get('declaration_honneur', 'NON')),
        "caution_provisoire": to_bool(merged.get('caution_provisoire', 'NON')),
        "note_moyens_humains": to_bool(merged.get('note_moyens_humains', 'NON')),
        "acheteur_detecte": merged.get('acheteur', '') or '',
        "depot_prospectus": merged.get('depot_prospectus', '') or '',
        "plan_charge": to_bool(merged.get('plan_charge', 'NON')),
        "moyens_humains_techniques": to_bool(merged.get('moyens_humains_techniques', 'NON')),
        "methodologie_travail": to_bool(merged.get('methodologie_travail', 'NON')),
        "memoire_technique": to_bool(merged.get('memoire_technique', 'NON')),
        "echantillon": to_bool(merged.get('echantillon', 'NON')),
        "acte_engagement": to_bool(merged.get('acte_engagement', 'NON')),
        "bordereau_prix": to_bool(merged.get('bordereau_prix', 'NON')),
        "attestations_ca": to_bool(merged.get('attestations_ca', 'NON')),
        "attestations_reference": to_bool(merged.get('attestations_reference', 'NON')),
        "rc_is_scanned": any_scanned,
    }
    try:
        supabase_client.table("tenders_3").upsert(data, on_conflict="reference").execute()
        return True
    except Exception: return False


def process_single_reference(reference: str) -> dict:
    rc_files = get_rc_files_from_supabase(reference)
    if not rc_files: return {"reference": reference, "status": "no_rc_files", "results": []}
    results = []; scanned_count = 0; groq_count = 0
    for rc in rc_files:
        try:
            r = process_rc_file(rc["filename"], rc["file_bytes"])
            if r.get('is_scanned'): scanned_count += 1
            if 'Groq' in str(r.get('method', '')): groq_count += 1
            results.append(r)
        except Exception: pass
    total_ok = sum(1 for r in results if "error" not in r and r.get('text_length', 0) > 0)
    saved = save_rc_extraction_to_supabase(reference, results)
    return {"reference": reference, "status": "saved" if saved else "extraction_error",
            "nb_files": len(rc_files), "nb_ok": total_ok, "nb_scanned": scanned_count, "nb_groq": groq_count, "results": results}


# ═══════════════════════════════════════════════════════════════
#  8. MAIN
# ═══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="RC Extraction v47 — OCR complet + Groq pour longs documents")
    parser.add_argument("--reference", "-r", type=str)
    parser.add_argument("--file", "-f", type=str)
    parser.add_argument("--all", "-a", action="store_true")
    parser.add_argument("--limit", "-l", type=int, default=None)
    args = parser.parse_args()

    groq_status = f"✅ {GROQ_MODEL}" if GROQ_API_KEY else "❌ Non configuré"
    db_status = "✅ Connecté" if supabase_client else "❌ Non connecté"

    print(f"\n{Fore.MAGENTA}{Style.BRIGHT}{'='*70}")
    print(f"{Fore.MAGENTA}{Style.BRIGHT}   {ICON_RC} RC EXTRACTION v47 FINAL")
    print(f"{Fore.MAGENTA}{Style.BRIGHT}   Groq: {groq_status} | Supabase: {db_status}")
    print(f"{Fore.MAGENTA}{Style.BRIGHT}   18 champs — OCR toutes pages + Groq chunking")
    print(f"{Fore.MAGENTA}{Style.BRIGHT}{'='*70}")

    if args.file:
        fp = Path(args.file)
        if not fp.exists(): print(f"\n{Fore.RED}❌ Fichier introuvable: {fp}"); return
        if fp.suffix.lower() not in ('.docx', '.doc', '.pdf'): print(f"\n{Fore.RED}❌ Format non supporté"); return
        with open(fp, 'rb') as f: fb = f.read()
        r = process_rc_file(fp.name, fb); display_result(r, 1, 1)
        return

    if args.all:
        references = get_all_references(limit=args.limit)
        if not references: print(f"\n{Fore.YELLOW}⚠️  Aucune référence trouvée"); return
        print(f"\n{ICON_BATCH} {Fore.CYAN}Traitement de {len(references)} référence(s)...\n")
        stats = {"total": len(references), "ok": 0, "no_rc": 0, "error": 0}
        for i, ref in enumerate(references, 1):
            print(f"  {ICON_BATCH} {Fore.WHITE}[{i}/{len(references)}] {ref}...", end=" ")
            result = process_single_reference(ref)
            if result['status'] == 'saved': print(f"{Fore.GREEN}✅ {result['nb_ok']} fichiers"); stats['ok'] += 1
            elif result['status'] == 'no_rc_files': print(f"{Fore.YELLOW}⚠️  Pas de RC"); stats['no_rc'] += 1
            else: print(f"{Fore.RED}❌ Erreur"); stats['error'] += 1
        print(f"\n{Fore.MAGENTA}{'='*70}")
        print(f"{Fore.MAGENTA}  📊 RÉSUMÉ BATCH")
        print(f"{Fore.MAGENTA}{'='*70}")
        print(f"  ✅ Succès : {Fore.GREEN}{stats['ok']}  ⚠️ Sans RC : {Fore.YELLOW}{stats['no_rc']}  ❌ Erreurs : {Fore.RED}{stats['error']}  📁 Total : {Fore.WHITE}{stats['total']}")
        print(f"{Fore.MAGENTA}{'='*70}\n")
        return

    if args.reference:
        rc_files = get_rc_files_from_supabase(args.reference)
        if not rc_files: return
        print(f"\n{Fore.CYAN}🚀 Traitement de {len(rc_files)} fichier(s) RC...")
        results = []
        scanned_count = 0
        groq_count = 0
        
        for i, rc in enumerate(rc_files, 1):
            try:
                r = process_rc_file(rc["filename"], rc["file_bytes"])
                if r.get('is_scanned'): scanned_count += 1
                if 'Groq' in str(r.get('method', '')): groq_count += 1
                results.append(r)
                display_result(r, i, len(rc_files))
            except Exception as e:
                print(f"\n  {ICON_CROSS} {Fore.RED}[{i}/{len(rc_files)}] {rc['filename']} — {e}")
                import traceback
                traceback.print_exc()
        
        total_ok = sum(1 for r in results if "error" not in r and r.get('text_length', 0) > 0)
        
        print(f"\n\n{Fore.MAGENTA}{'='*70}")
        print(f"{Fore.MAGENTA}  📊 RÉSUMÉ FINAL")
        print(f"{Fore.MAGENTA}{'='*70}")
        print(f"\n  Fichiers RC traités : {Fore.GREEN}{total_ok}/{len(rc_files)}")
        if scanned_count > 0: print(f"  {ICON_SCANNED} PDF RC scannés : {Fore.YELLOW}{scanned_count}")
        if groq_count > 0: print(f"  {ICON_AI} Groq utilisé : {Fore.CYAN}{groq_count}")
        
        if total_ok > 0:
            for icon, label, field in [
                (ICON_REF, "Attestations demandées", "attestations_demandees"),
                (ICON_TYPE_REF, "Types attestations", "types_attestations"),
                (ICON_NB_REF, "Nombre références", "nombre_references"),
                (ICON_CERTIF, "Certificat qualification", "classe_qualification"),
                (ICON_CA, "Chiffre d'affaires", "chiffre_affaires"),
                (ICON_DECL, "Déclaration honneur", "declaration_honneur"),
                (ICON_CAUT, "Caution provisoire", "caution_provisoire"),
                (ICON_NOTE, "Note moyens humains", "note_moyens_humains"),
                (ICON_PROSP, "Dépôt prospectus", "depot_prospectus"),
                (ICON_PLAN, "Plan de charge", "plan_charge"),
                (ICON_NOTE, "Moyens humains/techniques", "moyens_humains_techniques"),
                (ICON_METHOD, "Méthodologie travail", "methodologie_travail"),
                (ICON_MEMOIRE, "Mémoire technique", "memoire_technique"),
                (ICON_ECHANT, "Échantillon/prototype", "echantillon"),
                (ICON_ACTE, "Acte d'engagement", "acte_engagement"),
                (ICON_BORD, "Bordereau des prix", "bordereau_prix"),
                (ICON_ATT_CA, "Attestations CA", "attestations_ca"),
                (ICON_ATT_REF, "Attestations référence", "attestations_reference"),
            ]:
                count = sum(1 for r in results if "error" not in r and r.get('text_length', 0) > 0 and r.get(field))
                c = Fore.GREEN if count > 0 else Fore.RED
                bar = "█" * count + "░" * (total_ok - count)
                print(f"  {icon} {label:<35} {c}{count}/{total_ok}  {Fore.WHITE}{bar}")
        else:
            print(f"  {Fore.RED}Aucun fichier traité avec succès")
        
        if supabase_client and total_ok > 0:
            print(f"\n  {ICON_DB} {Fore.CYAN}Stockage automatique dans Supabase...")
            if save_rc_extraction_to_supabase(args.reference, results):
                print(f"  {ICON_CHECK} {Fore.GREEN}Résultats stockés avec succès !")
            else:
                print(f"  {ICON_CROSS} {Fore.RED}Échec du stockage.")
        
        print(f"\n{Fore.MAGENTA}{'='*70}\n")
        return

    print(f"\n{Fore.YELLOW}💡 Utilisation :")
    print(f"{Fore.YELLOW}   --reference REF  : Traiter une référence")
    print(f"{Fore.YELLOW}   --file FICHIER   : Traiter un fichier local")
    print(f"{Fore.YELLOW}   --all            : Traiter TOUTES les références")
    print(f"{Fore.YELLOW}   --all --limit 10 : Traiter 10 premières références")


if __name__ == "__main__":
    main()