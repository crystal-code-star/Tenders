"""
DOCX Extractor — Enhanced for Tender Intelligence
Mirrors the Excel extractor: full structural extraction (paragraphs, tables,
lists, headers/footers, styles, hyperlinks, comments, tracked changes,
sections) plus heuristic content classification (Pricing / Technical /
Administrative / Inventory / Calculation / Reference / General) so the
downstream LLM analyzer gets the same kind of rich, labeled context for Word
files that it already gets for Excel sheets.

Supports:
  - .docx / .docm  (via python-docx, full fidelity)
  - .doc            (legacy binary — requires LibreOffice/`soffice` on PATH
                      to convert to .docx first; raises a clear error if
                      unavailable, same pattern as the xlrd-missing case in
                      the Excel extractor)
"""

import io
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional
from fastapi import HTTPException

try:
    import docx
    from docx.document import Document as _DocumentObj
    from docx.oxml.ns import qn
    from docx.table import Table as _DocxTable
    from docx.text.paragraph import Paragraph as _DocxParagraph
except ImportError:
    docx = None


# ─── Heuristic classification (same keyword families as the Excel extractor) ───

_PRICING_KW = ["prix", "bpu", "bp", "dqe", "dp", "quantité", "unitaire", "total",
               "montant", "devise", "dh", "mad", "bordereau"]
_TECH_KW = ["équipement", "matériel", "specification", "technique", "données",
            "performance", "puissance", "débit", "pression", "tension"]
_ADMIN_KW = ["administratif", "conditions", "clauses", "attestation", "certificat",
             "qualification", "référence", "caution", "soumissionnaire", "candidature"]
_INV_KW = ["inventaire", "liste", "stock", "quantité", "désignation", "référence"]
_CALC_KW = ["calcul", "estimation", "budget", "prévision", "coût"]
_REF_KW = ["référence", "annexe", "document", "index"]


def _classify_text(text: str) -> str:
    """Same heuristic family as _classify_sheet() in the Excel extractor,
    applied to a chunk of paragraph/heading/table text instead of a sheet."""
    t = text.lower()
    scores = {
        "Pricing": sum(1 for kw in _PRICING_KW if kw in t),
        "Technical": sum(1 for kw in _TECH_KW if kw in t),
        "Administrative": sum(1 for kw in _ADMIN_KW if kw in t),
        "Inventory": sum(1 for kw in _INV_KW if kw in t),
        "Calculation": sum(1 for kw in _CALC_KW if kw in t),
        "Reference": sum(1 for kw in _REF_KW if kw in t),
    }
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "General"


# ─── Low-level formatting helpers ───────────────────────────────────

def _run_color(run) -> Optional[str]:
    try:
        color = run.font.color
        if color is None or color.rgb is None:
            return None
        return f"#{str(color.rgb)}"
    except Exception:
        return None


def _paragraph_alignment(paragraph) -> Optional[str]:
    try:
        align = paragraph.alignment
        return str(align) if align is not None else None
    except Exception:
        return None


def _list_info(paragraph) -> dict:
    """Detect numbering/bullet lists via style name and raw numPr XML,
    since python-docx doesn't expose list type/level directly."""
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    is_list_style = bool(re.search(r"list\s*(bullet|number|paragraph)", style_name, re.I))

    numPr = paragraph._p.find(qn("w:pPr") + "/" + qn("w:numPr")) if paragraph._p is not None else None
    ilvl = None
    numId = None
    if numPr is not None:
        ilvl_el = numPr.find(qn("w:ilvl"))
        numId_el = numPr.find(qn("w:numId"))
        if ilvl_el is not None:
            ilvl = ilvl_el.get(qn("w:val"))
        if numId_el is not None:
            numId = numId_el.get(qn("w:val"))

    is_list = is_list_style or numPr is not None
    list_type = None
    if is_list:
        if "bullet" in style_name.lower():
            list_type = "bullet"
        elif "number" in style_name.lower():
            list_type = "number"
        else:
            list_type = "unknown"

    return {
        "is_list": is_list,
        "list_type": list_type,
        "list_level": int(ilvl) if ilvl is not None else (0 if is_list else None),
        "num_id": numId,
    }


def _paragraph_hyperlinks(paragraph) -> list:
    """python-docx doesn't surface hyperlink text/target directly; parse the
    raw XML for w:hyperlink elements and resolve the target via the part's
    relationship map."""
    links = []
    try:
        for hl in paragraph._p.findall(qn("w:hyperlink")):
            rid = hl.get(qn("r:id"))
            text = "".join(node.text or "" for node in hl.iter(qn("w:t")))
            target = None
            if rid:
                try:
                    target = paragraph.part.rels[rid].target_ref
                except Exception:
                    target = None
            if text or target:
                links.append({"text": text, "target": target})
    except Exception:
        pass
    return links


def _paragraph_is_hidden(paragraph) -> bool:
    try:
        for run in paragraph.runs:
            if run.font.hidden:
                return True
    except Exception:
        pass
    return False


def _extract_runs(paragraph) -> list:
    runs = []
    for run in paragraph.runs:
        if run.text == "":
            continue
        runs.append({
            "text": run.text,
            "bold": bool(run.bold),
            "italic": bool(run.italic),
            "underline": bool(run.underline),
            "strike": bool(run.font.strike),
            "font_name": run.font.name,
            "font_size": run.font.size.pt if run.font.size else None,
            "color": _run_color(run),
            "highlight": str(run.font.highlight_color) if run.font.highlight_color else None,
        })
    return runs


def _heading_level(paragraph) -> Optional[int]:
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    m = re.match(r"Heading\s*(\d+)", style_name, re.I)
    if m:
        return int(m.group(1))
    if style_name.lower() in ("title",):
        return 0
    return None


def _tracked_changes(paragraph) -> dict:
    """Detect w:ins / w:del elements (Track Changes) inside a paragraph."""
    insertions, deletions = [], []
    try:
        for ins in paragraph._p.findall(qn("w:ins")):
            author = ins.get(qn("w:author"))
            text = "".join(t.text or "" for t in ins.iter(qn("w:t")))
            if text:
                insertions.append({"author": author, "text": text})
        for de in paragraph._p.findall(qn("w:del")):
            author = de.get(qn("w:author"))
            text = "".join(t.text or "" for t in de.iter(qn("w:delText")))
            if text:
                deletions.append({"author": author, "text": text})
    except Exception:
        pass
    return {"insertions": insertions, "deletions": deletions}


def _build_paragraph_dict(paragraph, index: int) -> dict:
    text = paragraph.text
    style_name = (paragraph.style.name or "Normal") if paragraph.style else "Normal"
    heading_level = _heading_level(paragraph)
    list_info = _list_info(paragraph)
    tracked = _tracked_changes(paragraph)

    return {
        "index": index,
        "text": text,
        "style": style_name,
        "heading_level": heading_level,
        "alignment": _paragraph_alignment(paragraph),
        "is_list": list_info["is_list"],
        "list_type": list_info["list_type"],
        "list_level": list_info["list_level"],
        "hyperlinks": _paragraph_hyperlinks(paragraph),
        "hidden": _paragraph_is_hidden(paragraph),
        "page_break_before": bool(paragraph.paragraph_format.page_break_before),
        "runs": _extract_runs(paragraph),
        "tracked_changes": tracked,
        "content_type": _classify_text(text) if text.strip() else "General",
    }


def _build_table_dict(table: "_DocxTable", index: int) -> dict:
    headers = []
    rows = []
    grid = []
    for r_idx, row in enumerate(table.rows):
        row_cells = []
        for c_idx, cell in enumerate(row.cells):
            cell_text = cell.text
            row_cells.append({
                "row": r_idx,
                "col": c_idx,
                "text": cell_text,
                "bold": any(r.bold for p in cell.paragraphs for r in p.runs if r.bold is not None),
            })
        grid.append(row_cells)
        if r_idx == 0:
            headers = [c["text"] for c in row_cells]
        else:
            rows.append([c["text"] for c in row_cells])

    sample_text = " ".join(headers) + " " + " ".join(" ".join(r) for r in rows[:10])
    return {
        "index": index,
        "headers": headers,
        "rows": rows,
        "grid": grid,
        "row_count": len(table.rows),
        "col_count": len(table.columns),
        "table_type": _classify_text(sample_text),
        "style": table.style.name if table.style else None,
    }


def _build_lists(paragraphs: list) -> list:
    """Group consecutive list-item paragraphs into logical lists, mirroring
    the shape the analyzer already expects: [{"items": [...]}]."""
    lists = []
    current_items = []
    for p in paragraphs:
        if p["is_list"] and p["text"].strip():
            current_items.append(p["text"])
        else:
            if current_items:
                lists.append({"items": current_items})
                current_items = []
    if current_items:
        lists.append({"items": current_items})
    return lists


def _extract_headers_footers(document) -> dict:
    headers, footers = [], []
    for section in document.sections:
        try:
            h_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
            if h_text:
                headers.append(h_text)
        except Exception:
            pass
        try:
            f_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
            if f_text:
                footers.append(f_text)
        except Exception:
            pass
    return {"headers": headers, "footers": footers}


def _extract_core_properties(document) -> dict:
    try:
        props = document.core_properties
        return {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
            "last_modified_by": props.last_modified_by or "",
            "revision": props.revision,
            "category": props.category or "",
            "comments": props.comments or "",
        }
    except Exception:
        return {}


def _extract_comments(document) -> list:
    """python-docx has no public comments API; parse the comments part
    directly if present (word/comments.xml)."""
    comments = []
    try:
        part = document.part
        comments_part = None
        for rel in part.rels.values():
            if "comments" in rel.reltype and rel.target_part is not None:
                comments_part = rel.target_part
                break
        if comments_part is None:
            return comments
        from lxml import etree
        root = etree.fromstring(comments_part.blob)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for c in root.findall("w:comment", ns):
            author = c.get(qn("w:author"))
            date = c.get(qn("w:date"))
            text = "".join(t.text or "" for t in c.iter(qn("w:t")))
            comments.append({"author": author, "date": date, "text": text})
    except Exception:
        pass
    return comments


def _iter_body_elements(document):
    """Walk the document body in document order, yielding ('paragraph', obj)
    or ('table', obj) tuples so tables and paragraphs stay interleaved in
    their original reading order (matters when a price table sits between
    two headings)."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("paragraph", _DocxParagraph(child, document))
        elif child.tag == qn("w:tbl"):
            yield ("table", _DocxTable(child, document))


def _convert_doc_to_docx(file_bytes: bytes, file_name: str) -> bytes:
    """Legacy .doc (binary) isn't readable by python-docx. Convert via
    LibreOffice headless if available on PATH — same fallback pattern as
    installing xlrd for legacy .xls in the Excel extractor, except this
    conversion needs an external binary rather than a pip package."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise HTTPException(
            422,
            "Legacy .doc files require LibreOffice (`soffice`) to convert to .docx, "
            "and it isn't available in this environment. Please convert the file to "
            ".docx first (e.g. via Word or LibreOffice) and re-upload."
        )
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / (Path(file_name).stem + ".doc")
        src.write_bytes(file_bytes)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, str(src)],
                check=True, capture_output=True, timeout=60,
            )
        except Exception as e:
            raise HTTPException(422, f"Failed to convert legacy .doc file: {e}")
        out = Path(tmp) / (Path(file_name).stem + ".docx")
        if not out.exists():
            raise HTTPException(422, "LibreOffice conversion did not produce a .docx output")
        return out.read_bytes()


def extract_docx_structured(file_bytes: bytes, file_name: str = "") -> dict:
    if docx is None:
        raise HTTPException(422, "python-docx not installed. Run: pip install python-docx")

    ext = Path(file_name).suffix.lower()

    if ext == ".doc":
        file_bytes = _convert_doc_to_docx(file_bytes, file_name)
    elif ext not in (".docx", ".docm"):
        raise HTTPException(400, "Unsupported Word format")

    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise HTTPException(422, f"Invalid Word file: {e}")

    paragraphs = []
    tables = []
    body_order = []  # ["p", "p", "t", "p", ...] to preserve reading order

    p_idx = 0
    t_idx = 0
    for kind, element in _iter_body_elements(document):
        if kind == "paragraph":
            pd = _build_paragraph_dict(element, p_idx)
            paragraphs.append(pd)
            body_order.append({"type": "paragraph", "ref": p_idx})
            p_idx += 1
        else:
            td = _build_table_dict(element, t_idx)
            tables.append(td)
            body_order.append({"type": "table", "ref": t_idx})
            t_idx += 1

    lists = _build_lists(paragraphs)
    hf = _extract_headers_footers(document)
    core_props = _extract_core_properties(document)
    comments = _extract_comments(document)

    total_bold = sum(1 for p in paragraphs for r in p["runs"] if r["bold"])
    total_tracked_changes = sum(
        len(p["tracked_changes"]["insertions"]) + len(p["tracked_changes"]["deletions"])
        for p in paragraphs
    )
    section_type_counts = {}
    for p in paragraphs:
        if p["text"].strip():
            section_type_counts[p["content_type"]] = section_type_counts.get(p["content_type"], 0) + 1

    stats = {
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "list_count": len(lists),
        "comment_count": len(comments),
        "bold_run_count": total_bold,
        "tracked_change_count": total_tracked_changes,
        "section_type_counts": section_type_counts,
        "has_headers_footers": bool(hf["headers"] or hf["footers"]),
    }

    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "lists": lists,
        "body_order": body_order,
        "headers": hf["headers"],
        "footers": hf["footers"],
        "comments": comments,
        "document_properties": core_props,
        "stats": stats,
        "type": "docx",
    }